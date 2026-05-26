from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "report_assets"
OUT_DOCX = DOCS_DIR / "SAINT_DATN_Report_ENG_Draft_Balanced.docx"
OUT_MD = DOCS_DIR / "SAINT_DATN_Report_ENG_Draft_Balanced.md"


TITLE = "BUILDING A DISTRIBUTED NETWORK SECURITY MANAGEMENT SYSTEM FOR EDUCATIONAL ENVIRONMENTS"
SHORT_TITLE = "SAINT: Distributed Network Security Management for Education"
STUDENT = "Bui Xuan Son"
STUDENT_ID = "20225586"
SUPERVISOR = "[Supervisor Name]"
PROGRAM = "[Program Name]"
EMAIL = "[student-email]@sis.hust.edu.vn"
DEPARTMENT = "Computer Engineering"
SCHOOL = "School of Information and Communications Technology"


@dataclass(frozen=True)
class Figure:
    number: str
    title: str
    filename: str


FIGURES = {
    "2.1": Figure("2.1", "General use case diagram of the SAINT system.", "figure_2_1_general_use_case.png"),
    "2.2": Figure("2.2", "Detailed use case diagram for whitelist and classroom policy management.", "figure_2_2_detailed_use_case.png"),
    "2.3": Figure("2.3", "Business process for preparing and enforcing a classroom network policy.", "figure_2_3_business_process.png"),
    "4.1": Figure("4.1", "Package-level architecture of the SAINT system.", "figure_4_1_system_architecture.png"),
    "4.2": Figure("4.2", "Detailed package design of the Windows Agent.", "figure_4_2_agent_package_design.png"),
    "4.3": Figure("4.3", "MongoDB collection relationship model.", "figure_4_3_database_model.png"),
    "4.4": Figure("4.4", "Sequence of agent registration, synchronization, and reporting.", "figure_4_4_agent_sequence.png"),
    "5.1": Figure("5.1", "Whitelist synchronization and enforcement flow.", "figure_5_1_whitelist_enforcement.png"),
    "5.2": Figure("5.2", "RBAC-based data filtering for dashboard operations.", "figure_5_2_rbac_filtering.png"),
}


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text))


def font(size: int):
    for name in ("arial.ttf", "calibri.ttf", "segoeui.ttf"):
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, xy, text: str, fill: str, outline: str = "#9a1c20") -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    f = font(30)
    words = text.split()
    lines: list[str] = []
    line = ""
    max_width = x2 - x1 - 34
    for word in words:
        candidate = f"{line} {word}".strip()
        if draw.textlength(candidate, font=f) <= max_width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    total_h = len(lines) * 36
    y = y1 + ((y2 - y1 - total_h) // 2)
    for line in lines:
        w = draw.textlength(line, font=f)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, fill="#111111", font=f)
        y += 36


def arrow(draw: ImageDraw.ImageDraw, start, end, color: str = "#555555") -> None:
    draw.line([start, end], fill=color, width=4)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        head = [(x2, y2), (x2 - sign * 18, y2 - 10), (x2 - sign * 18, y2 + 10)]
    else:
        sign = 1 if y2 > y1 else -1
        head = [(x2, y2), (x2 - 10, y2 - sign * 18), (x2 + 10, y2 - sign * 18)]
    draw.polygon(head, fill=color)


def diagram_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1600, 950), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1600, 76], fill="#b5121b")
    draw.text((48, 22), title, fill="#ffffff", font=font(34))
    return img, draw


def create_diagram_assets() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    img, draw = diagram_canvas("General Use Case Diagram")
    draw_box(draw, (80, 140, 320, 260), "Administrator", "#fff5f5")
    draw_box(draw, (80, 400, 320, 520), "Teacher", "#fff5f5")
    draw_box(draw, (80, 650, 320, 770), "Windows Agent", "#eef7ff", "#20639b")
    use_cases = [
        ("Manage users and API keys", 550, 115),
        ("Manage groups", 950, 115),
        ("Manage whitelist entries", 550, 330),
        ("Activate whitelist profiles", 950, 330),
        ("Register and heartbeat", 550, 555),
        ("Upload network logs", 950, 555),
        ("View scoped dashboard data", 750, 750),
    ]
    for text, x, y in use_cases:
        draw.ellipse((x, y, x + 290, y + 110), fill="#f7f7f7", outline="#9a1c20", width=3)
        draw.text((x + 22, y + 36), text, fill="#111111", font=font(23))
    for start, end in [
        ((320, 200), (550, 170)), ((320, 200), (950, 170)), ((320, 200), (550, 385)),
        ((320, 460), (550, 385)), ((320, 460), (950, 385)), ((320, 460), (750, 805)),
        ((320, 710), (550, 610)), ((320, 710), (950, 610)),
    ]:
        arrow(draw, start, end)
    img.save(ASSET_DIR / FIGURES["2.1"].filename)

    img, draw = diagram_canvas("Detailed Use Case: Whitelist and Policy")
    draw_box(draw, (60, 170, 340, 310), "Teacher / Admin", "#fff5f5")
    boxes = [
        ("Create or import entries", 520, 120),
        ("Validate domain/IP/URL", 950, 120),
        ("Create profile", 520, 330),
        ("Activate profile for group", 950, 330),
        ("Agent pulls effective policy", 520, 570),
        ("Firewall applies allow rules", 950, 570),
    ]
    for text, x, y in boxes:
        draw_box(draw, (x, y, x + 330, y + 130), text, "#f7f7f7")
    for start, end in [
        ((340, 240), (520, 185)), ((850, 185), (950, 185)), ((340, 240), (520, 395)),
        ((850, 395), (950, 395)), ((1115, 460), (685, 570)), ((850, 635), (950, 635)),
    ]:
        arrow(draw, start, end)
    draw.text((470, 810), "The effective policy is the result of global whitelist, group whitelist, active profile, and per-agent policy state.", fill="#333333", font=font(26))
    img.save(ASSET_DIR / FIGURES["2.2"].filename)

    img, draw = diagram_canvas("Business Process: Classroom Policy Operation")
    process = [
        ("Prepare lesson profile", 65),
        ("Assign group and agents", 335),
        ("Activate policy", 605),
        ("Agents synchronize", 875),
        ("Firewall enforces", 1145),
        ("Teacher reviews logs", 1415),
    ]
    for text, x in process:
        draw_box(draw, (x, 360, x + 190, 520), text, "#eef7ff", "#20639b")
    for (_, x1), (_, x2) in zip(process, process[1:]):
        arrow(draw, (x1 + 190, 440), (x2, 440))
    draw_box(draw, (230, 650, 560, 790), "Administrator maintains users API keys and global rules", "#fff5f5")
    draw_box(draw, (1020, 650, 1350, 790), "Audit trail records management operations", "#fff5f5")
    arrow(draw, (395, 650), (435, 520), "#9a1c20")
    arrow(draw, (1185, 650), (1240, 520), "#9a1c20")
    img.save(ASSET_DIR / FIGURES["2.3"].filename)

    img, draw = diagram_canvas("SAINT Package-Level Architecture")
    draw_box(draw, (70, 140, 410, 290), "Presentation Web Templates PySide6 GUI", "#fff5f5")
    draw_box(draw, (600, 130, 1000, 300), "Controller Layer Flask Blueprints Agent Controller", "#f7f7f7")
    draw_box(draw, (1180, 140, 1510, 290), "Service Layer Auth RBAC Whitelist Logs", "#fff8e6", "#b47a00")
    draw_box(draw, (270, 560, 610, 720), "Persistence Models MongoDB Collections", "#f2f5f2", "#3f7d20")
    draw_box(draw, (800, 535, 1170, 745), "Endpoint Runtime Firewall DNS Sniffer Token Config", "#eef7ff", "#20639b")
    arrow(draw, (410, 215), (600, 215))
    arrow(draw, (1000, 215), (1180, 215))
    arrow(draw, (1350, 290), (500, 560))
    arrow(draw, (800, 640), (610, 640))
    arrow(draw, (985, 535), (985, 300))
    img.save(ASSET_DIR / FIGURES["4.1"].filename)

    img, draw = diagram_canvas("Detailed Package Design: Windows Agent")
    draw_box(draw, (70, 130, 380, 270), "gui_qt views and components", "#fff5f5")
    draw_box(draw, (540, 130, 850, 270), "controllers AgentController", "#f7f7f7")
    draw_box(draw, (1010, 130, 1320, 270), "core lifecycle registry token", "#f7f7f7")
    draw_box(draw, (120, 540, 430, 690), "whitelist sync state resolver", "#eef7ff", "#20639b")
    draw_box(draw, (560, 540, 870, 690), "firewall manager rules", "#f2f5f2", "#3f7d20")
    draw_box(draw, (1000, 540, 1310, 690), "capture and logging sender", "#eef7ff", "#20639b")
    arrow(draw, (380, 200), (540, 200))
    arrow(draw, (850, 200), (1010, 200))
    arrow(draw, (1165, 270), (1165, 540))
    arrow(draw, (695, 270), (275, 540))
    arrow(draw, (695, 270), (715, 540))
    arrow(draw, (430, 615), (560, 615))
    arrow(draw, (870, 615), (1000, 615))
    draw.text((260, 805), "The GUI receives state through queued signals; network, firewall, and capture workers remain outside the UI thread.", fill="#333333", font=font(26))
    img.save(ASSET_DIR / FIGURES["4.2"].filename)

    img, draw = diagram_canvas("MongoDB Collection Relationship Model")
    collections = [
        ("users", 70, 140), ("groups", 430, 140), ("agents", 790, 140), ("logs", 1150, 140),
        ("api_keys", 70, 490), ("admin_sessions", 430, 490), ("whitelist", 790, 490), ("whitelist_profiles", 1150, 490),
        ("audit_logs", 250, 720), ("agent_policies", 970, 720),
    ]
    for text, x, y in collections:
        draw_box(draw, (x, y, x + 280, y + 115), text, "#f7f7f7")
    for start, end in [
        ((350, 197), (430, 197)), ((710, 197), (790, 197)), ((1070, 197), (1150, 197)),
        ((570, 255), (930, 490)), ((930, 255), (1110, 720)), ((930, 255), (930, 490)),
        ((570, 255), (1290, 490)), ((210, 605), (430, 545)), ((570, 605), (390, 720)),
    ]:
        arrow(draw, start, end)
    img.save(ASSET_DIR / FIGURES["4.3"].filename)

    img, draw = diagram_canvas("Agent Registration, Synchronization, and Reporting")
    actors = [("Agent", 160), ("Server API", 650), ("MongoDB", 1130)]
    for text, x in actors:
        draw_box(draw, (x, 120, x + 260, 235), text, "#f7f7f7")
        draw.line((x + 130, 235, x + 130, 830), fill="#999999", width=3)
    messages = [
        ("register with API key", 310, 700, 300),
        ("store/update agent", 780, 1180, 380),
        ("return agent token", 700, 310, 460),
        ("heartbeat + version", 310, 700, 540),
        ("agent-sync policy", 700, 310, 620),
        ("upload log batch", 310, 700, 700),
        ("persist logs", 780, 1180, 780),
    ]
    for text, x1, x2, y in messages:
        arrow(draw, (x1, y), (x2, y), "#20639b" if x2 > x1 else "#9a1c20")
        draw.text((min(x1, x2) + 25, y - 34), text, fill="#333333", font=font(23))
    img.save(ASSET_DIR / FIGURES["4.4"].filename)

    img, draw = diagram_canvas("Whitelist Sync and Enforcement")
    draw_box(draw, (70, 180, 420, 340), "Agent requests effective whitelist", "#eef7ff", "#20639b")
    draw_box(draw, (610, 180, 970, 340), "Server returns versions policy and entries", "#f7f7f7")
    draw_box(draw, (1160, 180, 1510, 340), "DNS resolver expands domains to IPs", "#fff8e6", "#b47a00")
    draw_box(draw, (340, 560, 700, 730), "Firewall manager snapshots and updates rules", "#f2f5f2", "#3f7d20")
    draw_box(draw, (900, 560, 1260, 730), "Default deny with explicit allow rules", "#fff5f5")
    arrow(draw, (420, 260), (610, 260))
    arrow(draw, (970, 260), (1160, 260))
    arrow(draw, (1335, 340), (1080, 560))
    arrow(draw, (610, 340), (520, 560))
    arrow(draw, (700, 645), (900, 645))
    draw.text((460, 810), "Self-allow, server allow, DNS allow, cleanup, and restore paths reduce operational risk.", fill="#333333", font=font(27))
    img.save(ASSET_DIR / FIGURES["5.1"].filename)

    img, draw = diagram_canvas("RBAC and Teacher Data Filtering")
    draw_box(draw, (90, 180, 420, 340), "Teacher Dashboard Request", "#fff5f5")
    draw_box(draw, (590, 180, 990, 340), "Authentication + Current User Context", "#f7f7f7")
    draw_box(draw, (1170, 180, 1510, 340), "RBAC Service", "#fff8e6", "#b47a00")
    draw_box(draw, (320, 580, 690, 740), "Group Ownership Filter", "#eef7ff", "#20639b")
    draw_box(draw, (910, 580, 1280, 740), "Scoped Agents Whitelist Logs Profiles", "#eef7ff", "#20639b")
    arrow(draw, (420, 260), (590, 260))
    arrow(draw, (990, 260), (1170, 260))
    arrow(draw, (1340, 340), (505, 580))
    arrow(draw, (690, 660), (910, 660))
    draw.text((400, 820), "Admin receives global scope; Teacher receives only groups assigned through group.teacher_ids.", fill="#333333", font=font(27))
    img.save(ASSET_DIR / FIGURES["5.2"].filename)


ACKNOWLEDGEMENTS = (
    "I would like to express my sincere gratitude to my supervisor, "
    f"{SUPERVISOR}, for the guidance, feedback, and technical orientation provided during this graduation "
    "thesis. I am also thankful to the lecturers of the School of Information and Communications Technology "
    "for the knowledge of computer networks, software engineering, databases, and information security that "
    "formed the foundation of this project. My appreciation goes to my family for their constant encouragement "
    "and to my classmates for discussions, testing suggestions, and practical comments during development. "
    "This thesis allowed me to combine distributed system design, network access control, role-based authorization, "
    "and Windows endpoint programming into a complete educational-lab management system. The experience helped "
    "me improve both implementation discipline and the ability to evaluate security risks in real deployment contexts."
)

ABSTRACT = (
    "Network access management in educational computer laboratories is difficult because many student machines must "
    "be controlled during class sessions while teachers still need flexible access to learning resources. Traditional "
    "approaches such as manual firewall configuration, proxy-only filtering, or standalone monitoring tools require "
    "considerable administrative effort and often lack a teacher-oriented workflow. This thesis follows a distributed "
    "client-server approach and builds SAINT, a centralized network security management system for educational "
    "environments. The system consists of a Flask and MongoDB server that provides REST APIs, a server-rendered web "
    "dashboard, SocketIO notifications, JWT authentication, API-key-based agent enrollment, audit logging, and "
    "role-based access control. Windows agents are implemented in Python with a PySide6 interface and background "
    "components for registration, heartbeat reporting, whitelist synchronization, DNS resolution, Windows Firewall "
    "rule management, packet capture with Scapy, domain extraction, encrypted configuration, and log delivery. The "
    "main contribution is a source-level implementation of whitelist-based network enforcement that can be managed "
    "centrally but applied locally on endpoint machines. The system also introduces versioned whitelist synchronization, "
    "teacher-scoped data filtering through RBAC, lesson-oriented whitelist profiles, and network log collection for "
    "post-session analysis. Static source inspection shows a modular server architecture with controller, service, and "
    "model layers, approximately seventy REST route declarations under the /api prefix, twelve MongoDB collections, "
    "and dedicated server tests for agents, authentication, groups, whitelist, logs, audit, and teacher data filtering. "
    "The result is a practical prototype that demonstrates how a laboratory network can be managed through a central "
    "dashboard while still applying enforcement on individual Windows endpoints."
)


def normal(text: str) -> dict:
    return {"type": "p", "text": text}


def heading(level: int, text: str) -> dict:
    return {"type": "h", "level": level, "text": text}


def table(number: str, title: str, headers: list[str], rows: list[list[str]]) -> dict:
    return {"type": "table", "number": number, "title": title, "headers": headers, "rows": rows}


def fig(number: str) -> dict:
    return {"type": "figure", "figure": FIGURES[number]}


def extract_server_api_rows(limit: int | None = None) -> list[list[str]]:
    source = ROOT / "report" / "server" / "04_TONG_HOP_API_VA_GIAI_THICH_SERVER.md"
    if not source.exists():
        return []
    rows: list[list[str]] = []
    for line in source.read_text(encoding="utf-8").splitlines():
        if not line.startswith("| ") or "---" in line or "Method" in line:
            continue
        cells = [cell.strip().replace("`", "") for cell in line.strip("|").split("|")]
        if len(cells) < 8:
            continue
        method, path, handler, auth, _purpose, _input, _output, service = cells[:8]
        auth = (
            auth.replace("Bearer access token và/hoặc refresh token", "Bearer access token and/or refresh token")
            .replace("Token trong JSON body hoặc Authorization header", "Token in JSON body or Authorization header")
            .replace("hoặc", "or")
            .replace("và", "and")
        )
        rows.append([method, path, handler, auth, service])
        if limit and len(rows) >= limit:
            break
    return rows


def python_inventory_rows() -> list[list[str]]:
    rows: list[list[str]] = []
    for area in ("agent", "server"):
        base = ROOT / area
        for py_file in sorted(base.rglob("*.py")):
            if any(part in {"__pycache__", ".pytest_cache"} for part in py_file.parts):
                continue
            if "dist" in py_file.parts or "build" in py_file.parts:
                continue
            relative = py_file.relative_to(ROOT).as_posix()
            try:
                tree = __import__("ast").parse(py_file.read_text(encoding="utf-8"))
            except Exception:
                rows.append([area, relative, "-", "-"])
                continue
            classes = [node.name for node in tree.body if isinstance(node, __import__("ast").ClassDef)]
            functions = [node.name for node in tree.body if isinstance(node, (__import__("ast").FunctionDef, __import__("ast").AsyncFunctionDef))]
            rows.append([
                area,
                relative,
                ", ".join(classes[:6]) + (" ..." if len(classes) > 6 else ""),
                ", ".join(functions[:8]) + (" ..." if len(functions) > 8 else ""),
            ])
    return rows


def add_paragraphs(elements: list[dict], paragraphs: list[str]) -> None:
    for paragraph in paragraphs:
        elements.append(normal(paragraph))


def long_use_case_tables() -> list[dict]:
    use_cases = [
        (
            "2.4",
            "Use case UC-01: Dashboard authentication and role initialization.",
            [
                ["Item", "Description"],
                ["Primary actor", "Administrator or Teacher."],
                ["Precondition", "A valid active account exists and the server is reachable."],
                ["Main flow", "The user submits credentials, the server validates the password hash, creates session or token context, loads role and permission information, and redirects the user to authorized dashboard pages."],
                ["Alternative flow", "If credentials are invalid or the account is locked, the server rejects the login and does not create an authenticated context."],
                ["Postcondition", "The server has an authenticated current-user context and subsequent requests can be checked by role, permission, and group ownership."],
            ],
        ),
        (
            "2.5",
            "Use case UC-02: Agent registration and heartbeat.",
            [
                ["Item", "Description"],
                ["Primary actor", "Windows Agent."],
                ["Precondition", "The agent has a server URL and a valid API key for enrollment."],
                ["Main flow", "The agent sends machine and runtime metadata to the registration endpoint, the server validates the API key, creates or updates the agent record, returns identity and token data, and later accepts heartbeat updates."],
                ["Alternative flow", "If the API key is missing, revoked, expired, or lacks registration permission, the registration attempt is rejected and no trusted agent identity is issued."],
                ["Postcondition", "The server can display the agent, update last-seen time, and decide whether synchronization is required."],
            ],
        ),
        (
            "2.6",
            "Use case UC-03: Whitelist and profile management.",
            [
                ["Item", "Description"],
                ["Primary actor", "Administrator or Teacher."],
                ["Precondition", "The user is authenticated and has permission for the target global or group scope."],
                ["Main flow", "The user creates, imports, updates, deletes, activates, or deactivates whitelist entries and profiles. The server validates input, updates MongoDB documents, and changes version metadata used by agents."],
                ["Alternative flow", "If a teacher targets an unassigned group, RBAC filtering prevents the change."],
                ["Postcondition", "The effective policy for affected agents changes and can be pulled through the synchronization endpoint."],
            ],
        ),
        (
            "2.7",
            "Use case UC-04: Endpoint whitelist enforcement.",
            [
                ["Item", "Description"],
                ["Primary actor", "Windows Agent."],
                ["Precondition", "The agent is registered, has a valid token, and is allowed to apply firewall policy on the local machine."],
                ["Main flow", "The agent requests the effective whitelist, resolves domains, creates self/server/DNS allow rules, snapshots current state, applies allow rules, and enables restrictive behavior when whitelist-only mode is active."],
                ["Alternative flow", "If synchronization or DNS resolution fails, the agent keeps the previous known policy or enters a safe state depending on configuration."],
                ["Postcondition", "Allowed destinations remain reachable and non-whitelisted traffic can be blocked by local firewall rules."],
            ],
        ),
        (
            "2.8",
            "Use case UC-05: Network log collection and dashboard review.",
            [
                ["Item", "Description"],
                ["Primary actor", "Windows Agent and dashboard user."],
                ["Precondition", "Packet capture is active and the agent can authenticate with the server."],
                ["Main flow", "The agent captures metadata, extracts domains from supported sources, batches records, uploads logs, and the dashboard user queries scoped logs and statistics."],
                ["Alternative flow", "If upload fails, logs remain queued until retry behavior sends them later."],
                ["Postcondition", "MongoDB stores network activity records and dashboard queries return data according to RBAC scope."],
            ],
        ),
    ]
    return [table(number, title, headers, rows) for number, title, (headers, *rows) in use_cases]


def insert_before_heading(elements: list[dict], heading_text: str, additions: list[dict]) -> None:
    for idx, element in enumerate(elements):
        if element["type"] == "h" and element["text"] == heading_text:
            elements[idx:idx] = additions
            return
    raise ValueError(f"Heading not found: {heading_text}")


def insert_after_heading(elements: list[dict], heading_text: str, additions: list[dict]) -> None:
    for idx, element in enumerate(elements):
        if element["type"] == "h" and element["text"] == heading_text:
            elements[idx + 1:idx + 1] = additions
            return
    raise ValueError(f"Heading not found: {heading_text}")


def remove_top_level_sections(elements: list[dict], section_titles: set[str]) -> list[dict]:
    result: list[dict] = []
    skip = False
    for element in elements:
        if element["type"] == "h" and element["level"] == 1:
            skip = element["text"] in section_titles
        if not skip:
            result.append(element)
    return result


def rebalance_main_chapters(elements: list[dict]) -> list[dict]:
    ch1_additions = [
        table(
            "1.2",
            "Traceability between objectives, scope, and thesis deliverables.",
            ["Thesis objective", "In-scope implementation", "Out-of-scope boundary"],
            [
                ["Centralized policy authority", "Flask server, REST APIs, MongoDB collections, dashboard pages, audit records", "Enterprise identity federation and multi-tenant cloud hosting"],
                ["Teacher-oriented classroom workflow", "Groups, teacher assignments, whitelist profiles, scoped logs", "Student-facing portal and attendance management"],
                ["Endpoint enforcement", "Windows agent, whitelist synchronization, DNS resolution, firewall manager", "Linux/macOS enforcement and router-level enforcement"],
                ["Monitoring and review", "Packet metadata capture, domain extraction, log upload, statistics, export", "Full payload inspection and HTTPS decryption"],
                ["Deployment prototype", "Docker-related server files and PyInstaller agent specification", "Commercial update infrastructure and production monitoring stack"],
            ],
        ),
        normal(
            "Table 1.2 clarifies the boundary of the thesis. The project is intentionally complete enough to demonstrate a distributed security-management workflow, but it does not claim to solve every endpoint-security problem. This boundary is important because a graduation-thesis report should not overstate the result. The implemented system is a working prototype with meaningful security functions, not a finished commercial product."
        ),
        normal(
            "The objective also includes source-level explainability. Every major claim in the report can be traced to a source area: server controllers and services for APIs, server models for collections, RBAC configuration for authorization, agent controller and lifecycle modules for runtime coordination, whitelist and firewall modules for enforcement, and capture/logging modules for monitoring. This traceability is necessary because the thesis is evaluated from both product behavior and implementation design."
        ),
        normal(
            "Another scope decision is that the report focuses on safe analysis and documentation. The agent executable and runtime firewall behavior are not executed during report generation. This decision does not reduce the value of the thesis; it reflects responsible handling of a tool that can intentionally change network access. Runtime demonstrations should be performed separately in a prepared test environment."
        ),
    ]
    insert_before_heading(elements, "1.3 Tentative solution", ch1_additions)

    ch2_additions = [
        heading(3, "2.3.1 Description of use case: Dashboard authentication"),
        normal(
            "This use case begins when an administrator or teacher opens the dashboard login page and submits credentials. The server receives the request, validates the user record, checks password hash state, creates the authenticated context, and prepares role or permission information for later requests. The main success condition is not simply reaching a web page. The important condition is that subsequent controller and service calls can identify the current user and apply RBAC consistently."
        ),
        normal(
            "The most important alternative flow occurs when credentials are invalid, the account is inactive, or the account is locked. In that case the server must reject the request without revealing sensitive information. A second alternative flow occurs when a token is expired or revoked. The system should require refresh or re-login instead of silently accepting stale identity. This use case supports all later dashboard functions."
        ),
        heading(3, "2.3.2 Description of use case: Agent enrollment and heartbeat"),
        normal(
            "Agent enrollment begins when a Windows endpoint has the server URL and an enrollment API key. The agent sends registration information such as device identity, host metadata, and runtime information. The server validates the API key, creates or updates an agent record, and returns identity and token material. After that, heartbeat requests keep the server informed about online state and synchronization needs."
        ),
        normal(
            "The main risk in this use case is accepting an untrusted endpoint. Therefore, the API key must be validated before issuing agent identity. The second risk is stale dashboard status. If heartbeat is not updated correctly, a teacher may believe that a machine is governed by the current policy when it is actually offline. This use case therefore links authentication, monitoring, and operational trust."
        ),
        heading(3, "2.3.3 Description of use case: Whitelist profile preparation"),
        normal(
            "A teacher or administrator prepares a whitelist profile before a class session. The profile contains entries needed for a lesson, such as documentation domains, repository hosts, package mirrors, online judge websites, or learning platforms. The user validates and saves the profile, then activates it for a group when the lesson begins. The server updates state so that affected agents can receive the new effective policy."
        ),
        normal(
            "The use case has several alternative flows. A malformed domain or IP address should be rejected. A teacher should be prevented from editing a group that is not assigned to that teacher. An activation request should not leave the group in an ambiguous profile state. These alternatives show why whitelist management is more complex than a simple text list."
        ),
        heading(3, "2.3.4 Description of use case: Endpoint enforcement"),
        normal(
            "Endpoint enforcement begins after the agent receives an effective policy. The agent resolves domain entries, prepares local allow rules, preserves access to the server and DNS, applies project-owned firewall rules, and reports status. When whitelist-only mode is active, the endpoint should allow the required destinations and restrict other outbound traffic. This use case is the point where a dashboard policy becomes a local security control."
        ),
        normal(
            "The alternative flows are operationally important. DNS resolution may fail, the server may be unreachable, the user may lack local administrator privilege, or firewall changes may fail. The agent must handle these cases visibly. A silent failure would be dangerous because the dashboard might show an intended policy while the endpoint is not enforcing it. A failed restore path would also be dangerous because it could leave the endpoint disconnected."
        ),
        heading(3, "2.3.5 Description of use case: Network log review"),
        normal(
            "Network log review begins when the agent captures metadata and uploads batches to the server. A dashboard user then lists logs, filters them, views statistics, exports records, or clears records according to permission. The use case helps teachers and administrators understand what happened during a class session. It also provides evidence for improving whitelist entries in later sessions."
        ),
        normal(
            "The main alternative flow is RBAC denial. A teacher should not review logs from groups outside the assigned scope. Another alternative is incomplete domain extraction. Some encrypted or privacy-preserving traffic may not expose a domain through DNS, HTTP Host, or TLS SNI. The system should still store useful metadata without pretending that every packet can be mapped to a human-readable domain."
        ),
    ]
    insert_before_heading(elements, "2.4 Non-functional requirement", ch2_additions)

    ch3_additions = [
        heading(2, "3.8 Technology trade-off summary"),
        normal(
            "The selected technologies form a pragmatic stack for a student-built distributed system. Flask is smaller than a full-stack framework and therefore makes the controller-service-model organization easier to inspect. MongoDB is flexible for logs and agent metadata. JWT and API keys provide separate identity flows for users and machines. PySide6 supports a native desktop agent interface. Windows Firewall integration uses a built-in enforcement point instead of introducing a custom kernel driver."
        ),
        normal(
            "The trade-off is that each technology also creates responsibilities. Flask does not automatically enforce application architecture; the project must maintain clean controllers and services. MongoDB does not enforce relational constraints in the same way as a relational database; services must protect relationships and ownership. JWTs require careful expiration and revocation handling. Windows Firewall changes require explicit recovery paths. Scapy-based capture can provide metadata but should not be treated as complete traffic understanding."
        ),
        normal(
            "The technology choices therefore match the thesis scope. They are strong enough to demonstrate the architecture and contributions, but not so heavy that the project becomes dominated by deployment infrastructure. The next chapter shows how these choices are applied to the actual design."
        ),
    ]
    insert_before_heading(elements, "CHAPTER 4. DESIGN, IMPLEMENTATION, AND EVALUATION", ch3_additions)

    ch4_api_additions = [
        normal(
            "The server API implementation can be summarized by functional group. Authentication routes protect user identity and token lifecycle. User routes manage administrator and teacher accounts. API-key routes manage enrollment credentials. Agent routes manage endpoint identity and runtime status. Group routes manage classroom organization. Whitelist routes manage policy entries and synchronization. Profile routes manage lesson-oriented policy variants. Log routes manage observation data. Audit routes manage administrative accountability."
        ),
        table(
            "4.6",
            "Main server API groups and implementation meaning.",
            ["API group", "Implementation meaning", "Design concern"],
            [
                ["Authentication", "Login, logout, refresh, verify, profile, password change", "Separate identity from authorization"],
                ["Users and audit", "Account management and management-action history", "Protect administrator-only operations"],
                ["API keys", "Enrollment credential lifecycle", "Allow revocation and avoid plaintext secrets"],
                ["Agents", "Registration, heartbeat, details, group, policy, statistics", "Keep endpoint state consistent"],
                ["Groups", "Classroom groups and teacher assignment", "Drive RBAC and policy scope"],
                ["Whitelist", "Entries, import, export, bulk operations, synchronization", "Compute deterministic effective policy"],
                ["Profiles", "Lesson-oriented profile create/update/activate/deactivate", "Support teacher workflow without rewriting base policy"],
                ["Logs", "Agent upload, list, clear, export, statistics", "Keep access role-scoped and privacy-aware"],
            ],
        ),
        normal(
            "This API grouping is more important for the thesis than the raw number of endpoints. It shows that the backend is structured around the actual educational workflow. The full endpoint catalogue is left in the appendix because a thesis chapter should explain why the API exists, while an appendix can list method and path details."
        ),
    ]
    insert_before_heading(elements, "4.4 Testing", ch4_api_additions)

    ch4_testing_additions = [
        normal(
            "For final evaluation, testing should be organized into scenarios rather than only unit-level route calls. The administrator setup scenario verifies users, groups, API keys, whitelist entries, and audit records. The teacher scenario verifies assigned-group visibility and profile operations. The agent scenario verifies enrollment, heartbeat, synchronization, and log upload. The enforcement scenario verifies firewall behavior only on a disposable or test endpoint with a prepared recovery plan."
        ),
        table(
            "4.7",
            "Suggested evaluation scenarios for the final defense.",
            ["Scenario", "Observed function", "Expected evidence"],
            [
                ["Administrator setup", "Users, groups, API keys, whitelist, audit", "Dashboard operations and audit records"],
                ["Teacher operation", "RBAC and group ownership", "Teacher sees only assigned groups and scoped logs"],
                ["Agent enrollment", "API-key registration and heartbeat", "Agent appears online in dashboard"],
                ["Dry-run synchronization", "Effective whitelist and version metadata", "Agent receives expected policy without enforcement"],
                ["Limited enforcement", "Windows Firewall allow/default-deny behavior", "Allowed resources work and blocked resources fail"],
                ["Log collection", "Packet metadata capture and upload", "Dashboard shows scoped log records and statistics"],
                ["Failure handling", "Token, API key, DNS, server, rollback behavior", "Errors are visible and recovery path is documented"],
            ],
        ),
        normal(
            "These scenarios also help separate safe tests from risky tests. Authentication, RBAC, API behavior, and dashboard visibility can be evaluated without changing firewall rules. Enforcement tests should be performed only after dry-run synchronization succeeds. This sequence reflects the operational risk of the agent and should be described during the defense."
        ),
    ]
    insert_before_heading(elements, "4.5 Deployment", ch4_testing_additions)

    ch4_deployment_additions = [
        normal(
            "A recommended deployment process has six stages. The first stage is configuration review, where secrets, server URL, TLS, database settings, and enrollment keys are prepared outside source control. The second stage is server-only testing with seeded users, groups, and whitelist entries. The third stage is agent dry-run testing, where communication is verified without firewall enforcement. The fourth stage is limited enforcement on one or two disposable endpoints. The fifth stage is a small classroom pilot. The final stage is broader rollout with monitoring and backup."
        ),
        table(
            "4.8",
            "Recommended deployment stages for safer rollout.",
            ["Stage", "Main activity", "Exit condition"],
            [
                ["Configuration review", "Prepare secrets, server URL, API keys, TLS, and database settings", "No real secret is stored in source and all required variables are documented"],
                ["Server-only test", "Validate dashboard workflows with test data", "Authentication, RBAC, whitelist, logs, groups, and audit behave as expected"],
                ["Agent dry-run", "Register endpoint and verify communication without enforcement", "Heartbeat, token refresh, sync, DNS resolution, and log upload work"],
                ["Limited enforcement", "Enable whitelist-only mode on test machines", "Server, DNS, and learning resources remain reachable and rollback works"],
                ["Classroom pilot", "Use a lesson profile with a small group", "Teacher can activate policy and review logs"],
                ["Broader rollout", "Expand to more machines with monitoring and backup", "Support team can operate and recover consistently"],
            ],
        ),
        normal(
            "The deployment plan is part of the design because endpoint enforcement has real operational impact. A web dashboard can often be fixed by redeploying the server, but an endpoint firewall mistake may require local recovery. Therefore, the thesis treats safety and rollout order as engineering requirements rather than afterthoughts."
        ),
    ]
    insert_before_heading(elements, "CHAPTER 5. SOLUTION AND CONTRIBUTION", ch4_deployment_additions)

    contribution_expansions = {
        "5.1 Whitelist-based endpoint firewall enforcement": [
            normal(
                "This contribution is more than calling a firewall command. The design must translate high-level educational intent into concrete endpoint behavior. A teacher thinks in terms of learning resources, but the endpoint enforces through program rules, IP addresses, DNS reachability, and outbound policy. The agent therefore acts as an adapter between classroom policy and operating-system enforcement."
            ),
            normal(
                "The enforcement contribution also requires failure reasoning. If a policy cannot be applied, the system should not pretend success. If a rule is applied but prevents server communication, the agent may lose the ability to recover remotely. For that reason, the implementation separates firewall management from GUI code and includes recovery-oriented behavior."
            ),
        ],
        "5.2 Versioned whitelist synchronization": [
            normal(
                "The versioning contribution addresses a practical scaling problem. In a laboratory with many endpoints, repeatedly sending and applying the same policy wastes server bandwidth, agent CPU time, and firewall update effort. Version metadata gives both sides a shared language for deciding whether anything changed."
            ),
            normal(
                "Versioning also improves explainability. When an agent reports its known version and the server returns a newer effective version, an administrator can reason about why synchronization occurred. This is useful for debugging profile activation, group changes, and global whitelist updates."
            ),
        ],
        "5.3 Teacher-scoped RBAC and group ownership filtering": [
            normal(
                "The RBAC contribution is important because it is enforced in backend logic. A dashboard can hide buttons, but hidden buttons are not a security boundary. The source design uses current-user context and service-level filtering so that API responses themselves are scoped. This is the correct place to enforce classroom ownership."
            ),
            normal(
                "The teacher role also improves usability. Teachers can focus on assigned groups without seeing unrelated agents and logs. This reduces accidental mistakes and makes the system more acceptable in a real educational environment where many classes may share the same infrastructure."
            ),
        ],
        "5.4 Lesson-oriented whitelist profiles": [
            normal(
                "Profiles turn whitelist management from a technical list-editing task into a teaching workflow. A teacher can prepare a profile before class, activate it during class, and reuse or modify it later. This maps better to lesson planning than editing a global allow list every time."
            ),
            normal(
                "The profile concept also creates a path to future features. Scheduling, exam mode, approval workflow, and temporary overrides can all be built on top of the profile layer. This shows that the contribution is an extensible design element rather than a one-off endpoint."
            ),
        ],
        "5.5 Packet metadata capture and domain extraction": [
            normal(
                "The monitoring contribution is deliberately metadata-oriented. It does not claim to decrypt content or inspect private payloads. Instead, it focuses on operationally useful information: which endpoint generated traffic, when it occurred, which address or domain was involved, and how it relates to classroom policy."
            ),
            normal(
                "This contribution helps close the feedback loop. If students cannot reach a required website, logs can reveal missing domains or resolution behavior. If irrelevant domains appear during a lesson, teachers can adjust future profiles. Monitoring therefore supports both troubleshooting and policy improvement."
            ),
        ],
    }
    for heading_text, additions in contribution_expansions.items():
        insert_after_heading(elements, heading_text, additions)
    insert_before_heading(elements, "CHAPTER 6. CONCLUSION AND FUTURE WORK", [
        heading(2, "5.6 Integrated contribution evaluation"),
        normal(
            "The five contributions should not be evaluated as isolated features. Their value comes from integration. Whitelist enforcement gives the system the ability to control traffic, but without versioned synchronization it would be difficult to update many endpoints safely. Versioned synchronization makes policy delivery efficient, but without RBAC the policy could be edited or inspected by the wrong teacher. RBAC protects classroom boundaries, but without whitelist profiles the teacher workflow would still be inconvenient. Profiles support lesson planning, but without packet monitoring there would be little feedback after a session. Monitoring closes the loop by showing whether policy and classroom activity matched the expected learning scenario."
        ),
        normal(
            "This integration also explains why the system is described as distributed network security management rather than simply a firewall agent. A firewall agent alone can block traffic, but it does not solve user roles, group ownership, policy preparation, audit trails, or dashboard review. A web dashboard alone can store policies, but it does not enforce them on endpoints. SAINT combines both sides and therefore addresses a broader laboratory-management problem."
        ),
        normal(
            "From an engineering perspective, the contributions also demonstrate separation of concerns. The server is responsible for authority, persistence, and authorization. The agent is responsible for local enforcement and observation. The dashboard is responsible for human workflow. MongoDB stores state. REST APIs connect components. This separation makes the implementation easier to inspect and provides clear directions for future replacement or improvement of individual modules."
        ),
        normal(
            "From a security perspective, the contributions show a balance between control and recovery. The system can restrict traffic through whitelist-only policy, but it also needs self-allow, server allow, DNS allow, cleanup, and restore behavior. The system can collect logs, but it also needs RBAC to prevent inappropriate visibility. The system can enroll agents, but it also needs API-key revocation and token lifecycle management. These paired controls are important for a thesis about practical security software."
        ),
        normal(
            "From an educational perspective, the contributions map technical functions to classroom needs. Teachers think in terms of classes, lessons, and required resources. Administrators think in terms of accounts, endpoints, policies, and audit. Agents operate in terms of tokens, versions, DNS results, firewall rules, and packet metadata. The project connects these viewpoints through groups, profiles, synchronization, logs, and role-scoped dashboards."
        ),
        table(
            "5.2",
            "Before-and-after evaluation of the main contributions.",
            ["Problem before contribution", "Contribution", "Improved result"],
            [
                ["Policy exists only as dashboard data", "Endpoint firewall enforcement", "Policy becomes enforceable on Windows machines"],
                ["Agents repeatedly fetch or apply unchanged policy", "Versioned synchronization", "Policy updates become more efficient and explainable"],
                ["Teachers could accidentally reach unrelated data", "Teacher-scoped RBAC", "Dashboard and API data are limited by group ownership"],
                ["Whitelist editing is repeated for every lesson", "Whitelist profiles", "Lesson-specific policy can be prepared and reused"],
                ["Traffic activity is hard to understand", "Packet metadata capture and domain extraction", "Teachers and administrators can review scoped activity logs"],
            ],
        ),
        normal(
            "The table shows that each contribution removes a specific weakness from the baseline problem. The baseline is a laboratory where rules are manual, visibility is weak, and teacher workflow is not represented in the system. The improved result is a source-inspectable prototype where policy, identity, enforcement, monitoring, and classroom scope are connected."
        ),
        table(
            "5.3",
            "Risk register connected to the main contributions.",
            ["Risk", "Contribution affected", "Mitigation in current design", "Future improvement"],
            [
                ["Default-deny rule blocks connectivity", "Firewall enforcement", "Self-allow, server allow, DNS allow, snapshot, restore, cleanup", "Dry-run mode and automatic rollback"],
                ["Teacher sees unrelated data", "RBAC filtering", "Backend group filtering", "More systematic authorization tests"],
                ["Whitelist misses required dependency", "Profiles and synchronization", "Import/export and profile reuse", "Policy simulation and dependency suggestions"],
                ["Logs grow too quickly", "Packet monitoring", "Indexes and scoped log queries", "Retention policy and aggregation pipeline"],
                ["Agent package is tampered with", "Endpoint enforcement", "Packaging specification", "Signed updates and integrity verification"],
            ],
        ),
        normal(
            "The risk register shows that each contribution has an associated engineering risk. This is expected in a security-management project. A strong thesis does not hide these risks; it explains how the current design reduces them and where future work should continue."
        ),
    ])

    return remove_top_level_sections(elements, {
        "E. API GROUP EXPLANATION",
        "F. OPERATIONAL SAFETY AND DEPLOYMENT PLAN",
        "G. LIMITATIONS AND RISK REGISTER",
        "H. EVALUATION PLAN FOR FINAL DEFENSE",
        "I. FINAL EDITING CHECKLIST",
    })


def build_elements() -> list[dict]:
    elements: list[dict] = []

    elements += [
        heading(1, "Chapter 1. Introduction"),
        heading(2, "1.1 Motivation"),
        normal(
            "Educational computer laboratories usually contain many shared machines that are used by different classes, "
            "teachers, and learning scenarios. During programming, networking, or digital-literacy sessions, students "
            "need Internet access to official documents and learning platforms, but unrestricted access can distract "
            "from the lesson, increase exposure to unsafe websites, and make classroom supervision difficult. In many "
            "laboratories, network control still depends on manual firewall rules, router-level configuration, or the "
            "availability of an IT administrator. These approaches do not fit lessons that change every hour and do "
            "not give teachers direct control over the resources that should be available for a specific class."
        ),
        normal(
            "The practical problem addressed in this thesis is therefore not only packet filtering. It is the combination "
            "of centralized policy management, local endpoint enforcement, classroom-oriented authorization, and useful "
            "monitoring. A teacher should be able to manage the whitelist for assigned groups without seeing unrelated "
            "laboratory data. An administrator should be able to enroll endpoint agents, manage API keys, review audit "
            "events, and observe system health. Each endpoint should receive only the effective policy it needs and apply "
            "that policy locally, even though the policy is managed centrally."
        ),
        normal(
            "The importance of this problem is increasing as educational activities depend more on online tools, remote "
            "documentation, cloud services, and learning platforms. A laboratory needs a control mechanism that is simple "
            "enough for classroom operation, auditable enough for administrators, and flexible enough to support different "
            "lessons. It should help teachers focus student activity during class while still preserving the technical "
            "visibility required by IT staff."
        ),
        heading(2, "1.2 Objectives and Scope"),
        normal(
            "The objective of the thesis is to design, implement, and evaluate a prototype distributed network security "
            "management system for educational environments. The system focuses on laboratory computers managed by a "
            "server dashboard and Windows endpoint agents. The scope includes account management, RBAC, group assignment, "
            "API-key-based agent registration, heartbeat monitoring, whitelist management, whitelist profiles, log "
            "collection, audit trails, and local network enforcement through the Windows Firewall."
        ),
        table(
            "1.1",
            "Thesis objectives and implementation scope.",
            ["Objective", "Implementation scope"],
            [
                ["Centralized management", "Provide server-side APIs and dashboards for agents, groups, whitelist, logs, users, API keys, and audit trails."],
                ["Teacher-oriented authorization", "Use RBAC and group ownership so teachers operate only on assigned groups while administrators retain global control."],
                ["Endpoint enforcement", "Use a Windows agent to synchronize whitelist data, resolve domains, and apply allow rules with a default-deny policy when enabled."],
                ["Monitoring and accountability", "Capture network metadata on the endpoint and send structured logs to the server for search, statistics, and export."],
                ["Deployable prototype", "Package the server with Docker-related files and the agent with a PyInstaller specification while keeping runtime configuration external."],
            ],
        ),
        normal(
            "The thesis does not attempt to replace enterprise-grade firewalls, endpoint detection platforms, or full web "
            "content inspection products. HTTPS content decryption, browser-level enforcement, device attestation, and "
            "cross-platform endpoint support are outside the current scope. The prototype is designed to demonstrate a "
            "manageable security architecture for computer laboratories, not to be a commercial network security appliance."
        ),
        heading(2, "1.3 Proposed Approach"),
        normal(
            "The proposed approach is a distributed client-server architecture. The server is the policy authority and "
            "stores users, groups, agents, whitelist entries, profiles, logs, sessions, API keys, and audit events in "
            "MongoDB. The web dashboard communicates with the server through HTTP and SocketIO. Endpoint agents register "
            "with an API key, receive JWT credentials, periodically send heartbeat messages, synchronize whitelist versions, "
            "and upload log batches. This separation allows policies to be edited centrally while enforcement remains close "
            "to the actual network stack of each Windows machine."
        ),
        normal(
            "A whitelist-based design was selected because educational sessions often have a known set of learning websites, "
            "online judges, repositories, or documentation pages. Instead of trying to classify all unwanted traffic, the "
            "agent can allow the domains and IP addresses required for a session and block other outbound traffic when "
            "whitelist-only mode is active. The codebase also includes safety mechanisms such as self-allow rules, server "
            "and DNS allow rules, firewall policy snapshots, restore paths, and cleanup logic. These mechanisms reduce, "
            "but do not eliminate, the risk that a wrong whitelist could interrupt connectivity."
        ),
        heading(2, "1.4 Thesis Organization"),
        normal(
            "The remainder of the thesis is organized into five chapters. Chapter 2 analyzes requirements and use cases. "
            "Chapter 3 presents the theoretical and technological background. Chapter 4 describes the architecture, design, "
            "implementation, testing, and deployment artifacts. Chapter 5 highlights the main technical contributions and "
            "their relationship to the current source code. Chapter 6 concludes the thesis and proposes future work."
        ),
    ]

    elements += [
        heading(1, "Chapter 2. Requirement Survey and Analysis"),
        normal(
            "This chapter surveys the application context and derives the functional and non-functional requirements for "
            "SAINT. It starts from classroom network management needs, identifies system actors, and defines the main use "
            "cases that the implemented prototype must support. The chapter also clarifies boundaries so that the later "
            "design chapter can be evaluated against explicit requirements."
        ),
        heading(2, "2.1 Context Survey"),
        normal(
            "A computer laboratory differs from a normal office network because access control is strongly tied to a lesson. "
            "A programming class may require documentation websites, package repositories, and online judges. A networking "
            "class may require different tools and remote resources. A test session may require a much smaller access list. "
            "The same machines can be reused across these scenarios, so static firewall rules or router-level rules are "
            "too rigid. The system must support quick policy changes without requiring teachers to work directly with low-level "
            "firewall syntax."
        ),
        table(
            "2.1",
            "Comparison of common approaches for laboratory network control.",
            ["Approach", "Strength", "Limitation in laboratory context"],
            [
                ["Manual endpoint firewall configuration", "Uses built-in operating-system capability", "Difficult to update consistently across many machines and risky when non-experts edit rules."],
                ["Router or gateway filtering", "Centralized at the network boundary", "Often lacks per-class teacher workflow and may not distinguish individual lab computers or lesson profiles."],
                ["Proxy-based filtering", "Can provide central logging and URL control", "Requires client configuration and may be bypassed or unsuitable for all protocols."],
                ["Standalone monitoring tool", "Observes traffic with minimal policy impact", "Does not enforce a whitelist and usually lacks RBAC-based classroom management."],
                ["SAINT distributed model", "Combines central policy management with endpoint enforcement", "Requires trusted Windows agents and careful firewall recovery procedures."],
            ],
        ),
        heading(2, "2.2 Actors"),
        normal(
            "The implemented source code is organized around three main operational actors. The administrator manages system-wide "
            "resources, security configuration, and audit visibility. The teacher manages educational groups assigned to that "
            "teacher and can work with whitelist profiles for lessons. The agent represents a Windows endpoint that registers "
            "with the server and applies the policy delivered by the server. Students are affected by endpoint enforcement, "
            "but they are not a direct authenticated actor in the current dashboard."
        ),
        table(
            "2.2",
            "Primary actors and responsibilities.",
            ["Actor", "Responsibilities", "Main source components"],
            [
                ["Administrator", "Manage users, groups, API keys, agents, whitelist data, audit logs, and global dashboard operations.", "server/controllers, server/services, server/models"],
                ["Teacher", "Manage assigned groups, view scoped agents/logs, and operate whitelist profiles related to teaching sessions.", "RBAC service, group service, whitelist profile service"],
                ["Windows Agent", "Register, authenticate, send heartbeat, synchronize whitelist, apply firewall rules, capture network metadata, and upload logs.", "agent/core, agent/services, agent/firewall, agent/capture"],
            ],
        ),
        heading(2, "2.3 Functional Requirements"),
        normal(
            "The first functional requirement is authentication and authorization. Dashboard users must log in and obtain a "
            "session context. Server endpoints must differentiate administrators from teachers and apply data filtering "
            "for group-scoped operations. Agent registration must use API keys and subsequent agent calls must use JWT-based "
            "authentication. The source code implements these flows through authentication controllers, JWT services, API-key "
            "services, RBAC configuration, and decorators applied to controller routes."
        ),
        normal(
            "The second requirement is agent lifecycle management. The server must accept agent registration, maintain agent "
            "identity, process heartbeat messages, and expose agent status to the dashboard. The agent must store configuration, "
            "manage tokens, recover from expired tokens when possible, and maintain worker components without blocking the GUI. "
            "This requirement is implemented by the agent controller, lifecycle module, heartbeat sender, token manager, and "
            "server-side agent controller and service layers."
        ),
        normal(
            "The third requirement is whitelist management. Administrators and teachers need to create, update, import, export, "
            "and inspect whitelist entries. Teachers must be limited to their assigned groups. The server must compute effective "
            "whitelist data for an agent and return version metadata so that the agent can avoid unnecessary updates. The agent "
            "must resolve domains into IP addresses and translate the resulting policy into Windows Firewall rules."
        ),
        normal(
            "The fourth requirement is monitoring and logging. The agent should capture network metadata and send batches of "
            "logs to the server. The dashboard should provide list, statistics, clearing, and export functions. Logs must also "
            "respect RBAC, meaning a teacher should not inspect unrelated group traffic. The current source implements this "
            "through the packet sniffer, domain extraction logic, log sender, log controller, log model, and log service."
        ),
        table(
            "2.3",
            "Major use cases derived from the current implementation.",
            ["Use case", "Description", "Success condition"],
            [
                ["UC-01 User authentication", "Admin or Teacher logs in, receives session/JWT context, and reaches permitted dashboard pages.", "Authenticated user context is available and route guards enforce role and permission checks."],
                ["UC-02 Agent enrollment", "Endpoint agent registers using an API key and receives identity/token material.", "Server creates or updates the agent record and returns credentials for future calls."],
                ["UC-03 Whitelist synchronization", "Agent requests effective whitelist data using version information.", "Server returns entries, policy mode, active profile, and version metadata; agent applies only needed updates."],
                ["UC-04 Firewall enforcement", "Agent resolves allowed domains and updates Windows Firewall rules.", "Allowed destinations are explicitly permitted and other outbound traffic is restricted when whitelist-only mode is enabled."],
                ["UC-05 Network log analysis", "Agent captures traffic metadata and uploads logs for dashboard inspection.", "Logs are stored in MongoDB and visible through role-scoped list, statistics, and export endpoints."],
            ],
        ),
        heading(2, "2.4 Non-Functional Requirements"),
        normal(
            "Security is the most important non-functional requirement. Credentials must not be hard-coded in the report or "
            "exposed from environment files. The server uses JWT and API keys; the agent stores configuration with encryption "
            "support; password storage uses hashing; and administrative actions can be written to an audit trail. RBAC must be "
            "applied consistently so that teacher visibility is narrower than administrator visibility."
        ),
        normal(
            "Reliability is also essential because endpoint enforcement can affect network access. The agent code includes "
            "default-deny firewall capability, but it must also allow the agent process, the server endpoint, and DNS services. "
            "The firewall manager therefore needs snapshot, restore, cleanup, and self-allow behavior. In operational use, the "
            "agent should be tested in an isolated environment before enabling whitelist-only mode on classroom machines."
        ),
        normal(
            "Maintainability is addressed through modular source structure. The server uses a controller-service-model style "
            "with blueprints under the /api prefix. The agent separates GUI, controller, lifecycle, network, firewall, capture, "
            "logging, and configuration modules. This separation makes it easier to test server functions independently and "
            "to reason about endpoint risk without mixing GUI updates with firewall logic."
        ),
        heading(2, "2.5 Chapter Conclusion"),
        normal(
            "This chapter identified the key requirements for a distributed laboratory network access-control system. The "
            "requirements show why SAINT must combine centralized dashboard management, endpoint enforcement, RBAC, whitelist "
            "versioning, monitoring, and careful recovery behavior. These requirements guide the technology choices and system "
            "design presented in the next chapters."
        ),
    ]

    elements += [
        heading(1, "Chapter 3. Theoretical Background and Technologies"),
        normal(
            "This chapter presents the background concepts and implementation technologies used by SAINT. The discussion focuses "
            "on distributed client-server design, REST APIs, real-time dashboard updates, MongoDB document persistence, JWT and "
            "API-key authentication, role-based access control, Windows Firewall management, packet capture, DNS, TLS Server Name "
            "Indication, and the GUI framework used by the endpoint agent."
        ),
        heading(2, "3.1 Client-Server Architecture and REST APIs"),
        normal(
            "SAINT follows a client-server architecture in which the server is responsible for authentication, authorization, "
            "policy storage, data aggregation, and dashboard presentation. Endpoint agents and browsers are clients of this "
            "server. REST APIs are used for resource-oriented operations such as registering agents, sending heartbeat messages, "
            "listing logs, updating groups, managing API keys, and synchronizing whitelist data. This style is appropriate because "
            "most operations are request-response interactions over HTTP and can be protected with standard authentication headers "
            "or cookies."
        ),
        normal(
            "The server implementation uses Flask, a lightweight Python web framework, with application factory initialization. "
            "Controller modules define blueprints and route handlers, service modules contain business logic, and model modules "
            "wrap MongoDB collections and indexes. This separation reduces coupling between HTTP routing and persistence details. "
            "The source analysis identified approximately seventy REST route declarations under the /api prefix, which demonstrates "
            "that most dashboard and agent operations are represented as explicit API endpoints."
        ),
        heading(2, "3.2 Real-Time Communication with SocketIO"),
        normal(
            "In addition to REST, the server uses Flask-SocketIO with gevent-related dependencies. SocketIO is suitable for dashboard "
            "updates because agent status, log statistics, or policy changes may need to be reflected without forcing the user to "
            "manually refresh a page. In the current prototype, SocketIO complements rather than replaces REST. Persistent data still "
            "flows through the model and service layers, while real-time events are used to notify the web interface about changes."
        ),
        heading(2, "3.3 MongoDB Document Persistence"),
        normal(
            "MongoDB is used because the data model contains documents that evolve naturally over time: agents may include host "
            "metadata and status, logs include packet attributes, whitelist entries include domain or IP fields, and sessions include "
            "token identifiers. MongoDB also supports indexes for common lookup fields such as agent identifiers, group identifiers, "
            "usernames, roles, timestamps, API-key hashes, and JWT identifiers. In SAINT, model classes are responsible for collection "
            "access and index setup."
        ),
        heading(2, "3.4 Authentication, JWT, API Keys, and RBAC"),
        normal(
            "Dashboard authentication and agent authentication have different requirements. A human dashboard user authenticates with "
            "credentials and receives a session context. An endpoint agent enrolls with an API key and then uses JWTs for protected "
            "server calls. JSON Web Tokens provide a compact representation of identity and claims [4], while API keys provide a "
            "controlled bootstrap mechanism for new endpoint registration. Token revocation and session models add operational control "
            "when a credential needs to be invalidated."
        ),
        normal(
            "Role-based access control, introduced in the classical RBAC model [5], is used to separate administrator and teacher "
            "capabilities. In SAINT, administrators have system-wide authority, while teachers are filtered by group ownership. This "
            "is important because logs, agents, groups, and whitelist profiles can contain information about different classes. A "
            "teacher should be able to manage a lesson without seeing unrelated data."
        ),
        heading(2, "3.5 Endpoint Network Enforcement"),
        normal(
            "The Windows agent applies network policy locally. Windows Firewall is used as the enforcement point because it is already "
            "available on the target operating system and can block or allow traffic at the endpoint. The agent translates server-side "
            "whitelist entries into rules and supports a default-deny mode where traffic is blocked unless it matches an allowed domain "
            "or IP address. This model is more intrusive than passive monitoring, so it requires careful handling of self-allow rules, "
            "DNS access, server connectivity, rule cleanup, and restoration."
        ),
        normal(
            "Domain-based allow rules require DNS resolution because firewall rules generally operate on IP addresses rather than "
            "abstract domain names. DNS behavior is defined in RFC 1035 [7], but modern websites may resolve to multiple IP addresses "
            "or change addresses over time. The agent therefore includes DNS resolver logic and whitelist synchronization so that the "
            "local firewall can be updated as policy and name resolution change."
        ),
        heading(2, "3.6 Packet Capture and Domain Extraction"),
        normal(
            "Packet capture is used for monitoring and accountability. The agent uses Scapy-related functionality to inspect packet "
            "metadata. Domain extraction can use DNS queries, HTTP Host headers, and TLS Server Name Indication. DNS provides explicit "
            "name lookup information [7]. HTTP/1.1 defines the Host header used by clients to identify the target virtual host [10]. "
            "TLS extensions define Server Name Indication, which lets a client indicate the server name during TLS negotiation [9]. "
            "These metadata sources allow useful logging without decrypting application content."
        ),
        heading(2, "3.7 User Interface Technology"),
        normal(
            "The current agent GUI uses PySide6, the official Qt for Python binding. The GUI is separated from background workers "
            "through a controller and signal bridge. Worker components emit events into a queue, and the GUI thread drains the queue "
            "through a timer so that interface updates do not block packet capture, heartbeat, or synchronization. This design follows "
            "a model-view-presenter-like separation in which the view layer is not responsible for direct firewall or network logic."
        ),
        table(
            "3.1",
            "Main technologies used in the current source code.",
            ["Layer", "Technology", "Purpose"],
            [
                ["Server web/API", "Flask, Flask-CORS, Flask-SocketIO, gevent", "REST APIs, server-rendered dashboard, CORS, and real-time events."],
                ["Persistence", "MongoDB, PyMongo, Pydantic", "Document storage, indexes, validation-oriented structures, and collection access."],
                ["Security", "PyJWT, bcrypt, API keys, RBAC service", "User sessions, token issuance, password hashing, agent enrollment, and role filtering."],
                ["Agent GUI", "PySide6", "Windows desktop interface for status, settings, firewall, whitelist, and logs."],
                ["Agent networking", "requests, urllib3, dnspython, aiodns, netifaces", "Server communication, DNS resolution, and network interface information."],
                ["Agent enforcement/monitoring", "pywin32, Scapy, psutil, cryptography", "Windows integration, packet capture, process/system data, and encrypted configuration."],
            ],
        ),
        heading(2, "3.8 Chapter Conclusion"),
        normal(
            "This chapter explained the main concepts behind SAINT and mapped them to implementation technologies. Flask, MongoDB, "
            "JWT, API keys, RBAC, SocketIO, Windows Firewall, DNS, TLS SNI, Scapy, and PySide6 together provide the technical basis "
            "for a centralized yet endpoint-enforced network management prototype. The next chapter shows how these technologies are "
            "assembled into the actual system architecture."
        ),
    ]

    elements += [
        heading(1, "Chapter 4. Design, Implementation, Testing, and Deployment"),
        normal(
            "This chapter describes the design and implementation of SAINT based on the current source code. It presents the overall "
            "architecture, server design, agent design, database collections, API groups, dashboard behavior, testing scope, and deployment "
            "artifacts. All descriptions in this chapter are derived from static source analysis; no agent executable, server runtime, Docker "
            "service, packet capture, netsh command, or firewall-changing code was executed during report preparation."
        ),
        heading(2, "4.1 Overall Architecture"),
        normal(
            "The system contains four major parts: a web dashboard for administrators and teachers, a Flask server, MongoDB storage, "
            "and Windows endpoint agents. Figure 4.1 shows the architecture. The browser communicates with the server through HTTP and "
            "SocketIO. The server stores data in MongoDB and exposes REST APIs under the /api prefix. Each agent communicates with the "
            "server through authenticated HTTP requests and uses local components for whitelist synchronization, firewall enforcement, "
            "packet capture, and log upload."
        ),
        fig("4.1"),
        normal(
            "This architecture separates policy authority from policy enforcement. The server decides what data a user or agent can access "
            "and stores the authoritative whitelist state. The agent applies the resulting effective policy on the endpoint machine. This "
            "is useful for laboratories because a central administrator can manage global settings while each endpoint remains responsible "
            "for enforcing rules on its own network stack."
        ),
        heading(2, "4.2 Server Architecture"),
        normal(
            "The server follows a controller-service-model structure. Controllers receive HTTP requests, decode inputs, apply decorators, "
            "and call services. Services implement business rules such as registering agents, creating users, filtering teacher data, "
            "activating whitelist profiles, validating API keys, or recording audit logs. Models encapsulate MongoDB collections, document "
            "operations, and index creation. The application factory registers controller blueprints with the /api prefix and also provides "
            "server-rendered dashboard routes for pages such as agents, groups, whitelist, logs, API keys, user administration, audit, login, "
            "profile, and password change."
        ),
        table(
            "4.1",
            "MongoDB collections identified in the current source.",
            ["Collection", "Purpose"],
            [
                ["agents", "Endpoint identity, hostname, group assignment, status, heartbeat, and metadata."],
                ["agent_policies", "Runtime policy overrides or per-agent policy state."],
                ["api_keys", "Agent enrollment keys, hashes, permissions, status, and expiry information."],
                ["audit_logs", "Administrative and security-relevant action history."],
                ["groups", "Classroom or laboratory groups and teacher assignments."],
                ["logs", "Network access logs uploaded by agents."],
                ["admin_sessions", "Dashboard sessions and token identifiers."],
                ["users", "Administrator and teacher accounts, roles, password hashes, and lock state."],
                ["whitelist", "Global or group-specific domain, IP, URL, category, and active-state entries."],
                ["whitelist_meta", "Global whitelist version metadata."],
                ["whitelist_profiles", "Lesson-oriented whitelist profiles for groups and teachers."],
                ["revoked_tokens", "Revoked JWT identifiers and TTL cleanup data."],
            ],
        ),
        normal(
            "The API surface is broad because the dashboard and the agent are both clients of the server. Table 4.2 summarizes the main "
            "API groups. The full route list in the source-based technical report contains approximately seventy method-path entries, all "
            "registered under the /api prefix except the dashboard template routes."
        ),
        table(
            "4.2",
            "Main server API groups.",
            ["API group", "Representative paths", "Purpose"],
            [
                ["Authentication", "/api/admin/auth/login, /api/auth/refresh, /api/auth/verify", "Authenticate dashboard users, refresh tokens, verify token state, and support logout."],
                ["Users and audit", "/api/admin/users, /api/admin/audit", "Manage dashboard accounts and inspect administrative activity."],
                ["Agents", "/api/agents/register, /api/agents/heartbeat, /api/agents/<id>/policy", "Enroll endpoints, update status, retrieve details, and manage agent policies."],
                ["Groups and profiles", "/api/groups, /api/groups/<id>/profiles, /api/my-profiles", "Manage classroom groups, teacher assignment, and lesson-specific whitelist profiles."],
                ["Whitelist", "/api/whitelist, /api/whitelist/agent-sync, /api/whitelist/import", "Create, update, import, export, and synchronize effective whitelist data."],
                ["Logs", "/api/logs, /api/logs/stats, /api/logs/export", "Receive agent log batches and provide role-scoped dashboard inspection."],
                ["API keys", "/api/api-keys, /api/api-keys/validate, /api/api-keys/<id>/revoke", "Manage enrollment credentials for endpoint agents."],
            ],
        ),
        heading(2, "4.3 Agent Architecture"),
        normal(
            "The Windows agent is composed of a PySide6 GUI, a central controller, a signal bridge, lifecycle initialization, and several "
            "background components. The GUI displays status, firewall rules, IP whitelist data, logs, and settings. The controller is "
            "framework-independent and manages start/stop behavior, worker threads, statistics, and signal emission. The signal bridge "
            "drains queued events on the Qt main thread so that background work does not directly update widgets."
        ),
        normal(
            "Figure 4.2 shows the main lifecycle. On startup, the agent loads configuration, initializes components, registers or refreshes "
            "credentials, synchronizes whitelist data, and starts periodic heartbeat and log delivery. Packet capture and firewall enforcement "
            "operate as background capabilities. This architecture is important because packet monitoring and server communication must not "
            "freeze the user interface."
        ),
        fig("4.2"),
        heading(2, "4.4 Detailed Design of Policy Synchronization"),
        normal(
            "Policy synchronization is version-aware. The server keeps whitelist and group-related state, while the agent sends its known "
            "version information to the synchronization endpoint. The server can return effective whitelist entries, policy mode, active "
            "profile information, and version metadata. This reduces unnecessary data transfer and makes it possible for a class-specific "
            "profile to change without requiring manual intervention on each endpoint."
        ),
        normal(
            "The agent must convert high-level whitelist entries into enforceable local rules. Domain entries require DNS resolution. IP entries "
            "can be used more directly. When whitelist-only mode is active, the agent sets explicit allow rules and relies on default-deny behavior "
            "for other outbound traffic. Because this behavior can interrupt network connectivity, the implementation includes risk-reduction logic "
            "such as self-allow, server allow, DNS allow, snapshot, restore, and cleanup routines."
        ),
        heading(2, "4.5 Authentication and RBAC Design"),
        normal(
            "Authentication is split into human-user and agent flows. Human users log in through dashboard authentication routes. Agent enrollment "
            "uses API keys, and subsequent protected calls use JWT tokens. RBAC applies role and permission checks to dashboard operations. For "
            "teacher users, services apply group ownership filters so that queries for groups, agents, logs, and whitelist-related data return only "
            "the assigned scope. Figure 5.2 later illustrates this filtering path as a contribution."
        ),
        heading(2, "4.6 Testing"),
        normal(
            "The source tree contains server-side tests that cover the most important backend behavior. The tests indicate that the project was not "
            "implemented only as a user-interface prototype. Instead, agent registration, heartbeat handling, whitelist/log operations, authentication, "
            "teacher filtering, groups, and audit behavior are all represented in test files."
        ),
        table(
            "4.3",
            "Testing scope identified from server test files.",
            ["Test file or group", "Scope"],
            [
                ["test_agents.py and test_agent_full.py", "Agent registration, heartbeat, and agent API behavior."],
                ["test_whitelist_and_logs.py", "Whitelist management, synchronization, log receiving, and log queries."],
                ["test_users_auth.py", "User management and authentication behavior."],
                ["test_teacher_data_filtering.py", "RBAC behavior that limits teachers to assigned data."],
                ["test_groups.py", "Group creation, update, deletion, and assignment behavior."],
                ["test_audit.py", "Audit logging behavior for administrative operations."],
            ],
        ),
        normal(
            "No tests were executed while preparing this report because the request was to write the report and avoid actions that may start the "
            "server or affect networking. The testing description is therefore based on static inspection of the test source files and should be "
            "validated by running the test suite in a controlled development environment when preparing the final submission."
        ),
        heading(2, "4.7 Deployment Artifacts"),
        normal(
            "The server folder contains a Dockerfile, docker-compose file, requirements file, and .env-example. These artifacts indicate an intended "
            "containerized deployment path, while environment-specific secrets are kept outside source control. The agent folder contains a PyInstaller "
            "specification for building a Windows executable. In a real laboratory rollout, the server should be configured with production-grade secrets, "
            "TLS termination, MongoDB backup procedures, logging retention, and administrator account hardening. Endpoint rollout should be staged and "
            "tested on non-critical machines before enabling whitelist-only enforcement."
        ),
        heading(2, "4.8 Chapter Conclusion"),
        normal(
            "This chapter presented the architecture and implementation of SAINT. The design uses a Flask and MongoDB server for central control, a "
            "PySide6 Windows agent for endpoint enforcement, REST and SocketIO for communication, and RBAC to separate administrator and teacher "
            "responsibilities. The static source analysis shows that the implementation includes both backend management features and endpoint-level "
            "network functions."
        ),
    ]

    elements += [
        heading(1, "Chapter 5. Solution and Contribution"),
        normal(
            "This chapter highlights the main technical contributions of the thesis. The contributions are evaluated from the current source code rather "
            "than from older project documents. They include whitelist-based firewall enforcement, versioned whitelist synchronization, teacher-scoped RBAC, "
            "whitelist profiles, and packet monitoring with domain extraction."
        ),
        heading(2, "5.1 Whitelist-Based Firewall Enforcement"),
        normal(
            "The first contribution is endpoint-level whitelist enforcement. Instead of only displaying a list of allowed websites on the dashboard, the "
            "agent translates effective policy into Windows Firewall behavior. Figure 5.1 shows the flow from synchronization to DNS resolution and rule "
            "application. The main idea is that allowed destinations are created explicitly, while other outbound traffic can be blocked under whitelist-only "
            "mode."
        ),
        fig("5.1"),
        normal(
            "This contribution is significant because the enforcement decision happens on each endpoint. If the network contains multiple laboratory machines, "
            "each agent can independently enforce the effective policy received from the server. The server remains the management authority, but the endpoint "
            "does not require the network gateway to understand every classroom policy."
        ),
        normal(
            "The approach also introduces risk. A wrong default-deny rule can cut network access, including the agent's own connection to the server. The source "
            "therefore includes mitigation ideas such as self-allow rules, server URL allow rules, DNS allow rules, rule snapshots, cleanup, and restoration. "
            "These mechanisms are necessary for a security tool that modifies local firewall behavior."
        ),
        heading(2, "5.2 Versioned Whitelist Synchronization"),
        normal(
            "The second contribution is version-aware synchronization. The agent does not need to download and reapply the full policy on every heartbeat if "
            "nothing changed. Instead, it can compare known global and group versions with the server state and receive only the information needed to update "
            "local enforcement. This design reduces redundant work and supports a smoother classroom workflow."
        ),
        normal(
            "Versioned synchronization also helps separate global policy from group policy. An administrator may define global domains that are always required, "
            "while a teacher may activate a group-specific profile for a lesson. The agent receives the effective result, not every administrative detail. This "
            "keeps endpoint logic simpler and supports future extensions such as scheduled profiles or temporary exam modes."
        ),
        heading(2, "5.3 Teacher-Scoped RBAC"),
        normal(
            "The third contribution is applying RBAC to classroom ownership rather than only to menu visibility. Teacher accounts are not simply less powerful "
            "administrators. They receive filtered data according to group assignment. This means that queries for groups, agents, logs, and whitelist profiles "
            "must be scoped by the teacher's permitted groups. Figure 5.2 illustrates the authorization path."
        ),
        fig("5.2"),
        normal(
            "This feature matters in educational environments because one laboratory system can serve many classes. Without scoped filtering, a teacher could "
            "accidentally inspect or modify resources belonging to another class. The source code addresses this with RBAC configuration, current-user context, "
            "group teacher identifiers, and services that apply permission or ownership filters."
        ),
        heading(2, "5.4 Whitelist Profiles"),
        normal(
            "The fourth contribution is the whitelist profile concept. A profile represents a lesson-oriented or scenario-oriented whitelist set. Instead of "
            "editing the base whitelist for every class, a teacher can prepare or activate a profile for a specific group. This design is closer to classroom "
            "practice, where different activities require different allowed resources."
        ),
        normal(
            "Whitelist profiles also improve maintainability. A profile can be created, updated, activated, or deactivated through dedicated server routes. The "
            "server can then expose the active profile as part of the effective policy delivered to agents. Although the current prototype is not a full scheduling "
            "system, the profile layer provides a foundation for time-based activation in future work."
        ),
        heading(2, "5.5 Packet Monitoring and Domain Extraction"),
        normal(
            "The fifth contribution is network metadata collection on the endpoint. The agent captures packets and attempts to extract meaningful domain information "
            "from DNS, HTTP Host headers, and TLS SNI where available. This allows the server to show access logs and statistics that are more useful than raw IP "
            "addresses alone. The implementation avoids application-content decryption and focuses on metadata relevant to classroom supervision."
        ),
        normal(
            "This monitoring function complements enforcement. Firewall rules decide whether traffic should be allowed, while packet metadata helps teachers or "
            "administrators understand what happened during a session. The log sender batches data and the server stores logs in MongoDB, where the dashboard can "
            "list, filter, clear, export, or summarize them according to role-based visibility."
        ),
        table(
            "5.1",
            "Contribution traceability to source-level modules.",
            ["Contribution", "Representative source areas"],
            [
                ["Whitelist-based endpoint enforcement", "agent/firewall, agent/whitelist, agent/network"],
                ["Versioned synchronization", "server whitelist controller/service/model and agent whitelist manager"],
                ["Teacher-scoped RBAC", "server RBAC configuration, group service, log service, whitelist service"],
                ["Whitelist profiles", "server whitelist profile controller, service, and model"],
                ["Packet monitoring and domain extraction", "agent/capture and agent/logging_module"],
            ],
        ),
        heading(2, "5.6 Chapter Conclusion"),
        normal(
            "The main value of SAINT is the integration of these contributions into a single prototype. Centralized management alone would not enforce policy on "
            "student machines, and endpoint enforcement alone would not give teachers a manageable classroom workflow. By combining both sides with RBAC, versioned "
            "sync, profiles, monitoring, and auditability, the system demonstrates a practical direction for educational laboratory network control."
        ),
    ]

    elements += [
        heading(1, "Chapter 6. Conclusion and Future Work"),
        normal(
            "This chapter summarizes the thesis outcome and proposes future development directions. The chapter does not introduce new implementation details; it "
            "evaluates what has been achieved by the current prototype and what should be improved before broader deployment."
        ),
        heading(2, "6.1 Conclusion"),
        normal(
            "The thesis designed and implemented SAINT, a distributed network security management system for educational computer laboratories. The system combines "
            "a Flask and MongoDB server, a web dashboard, RBAC, JWT and API-key authentication, audit logging, whitelist management, SocketIO-supported updates, and "
            "a Windows endpoint agent. The agent registers with the server, manages token state, synchronizes whitelist policy, resolves domains, applies firewall "
            "rules, captures network metadata, and uploads logs."
        ),
        normal(
            "The implementation demonstrates that classroom network control can be modeled as a combination of central policy management and local endpoint enforcement. "
            "Administrators retain global control, teachers receive scoped access to their groups, and agents apply the effective whitelist on student machines. The "
            "source code also shows attention to operational risks through firewall recovery behavior, encrypted configuration support, token revocation, and audit trails."
        ),
        normal(
            "The prototype remains a graduation-thesis system rather than a production-ready security product. It should be evaluated more extensively under realistic "
            "laboratory load, with controlled network-failure scenarios, malformed whitelist entries, expired certificates, DNS changes, and accidental policy mistakes. "
            "Nevertheless, the current result provides a clear foundation for a deployable educational-lab network management system."
        ),
        heading(2, "6.2 Future Work"),
        normal(
            "Future work should first improve operational safety. The agent should provide a stronger emergency restore mechanism, explicit dry-run mode, better policy "
            "validation before firewall changes, and clearer recovery instructions for administrators. A staged rollout process should be added so that a new policy can "
            "be tested on one endpoint before applying it to an entire laboratory."
        ),
        normal(
            "The second direction is policy expressiveness. Whitelist profiles can be extended with schedules, exam modes, temporary overrides, import validation, and "
            "conflict detection between global and group rules. The server could also provide policy simulation so teachers know which domains and IP addresses will be "
            "allowed before the policy reaches endpoints."
        ),
        normal(
            "The third direction is monitoring and analytics. The dashboard can add richer traffic visualization, anomaly detection, per-session reports, and retention "
            "policies. Since the current system focuses on metadata, future versions should clearly document privacy boundaries and provide configurable data minimization "
            "for educational compliance."
        ),
        normal(
            "The final direction is deployment hardening. Production use should include TLS everywhere, stronger secret management, signed agent updates, service-mode "
            "operation, centralized backup, observability, load testing, and cross-version migration scripts. With these improvements, SAINT could evolve from a prototype "
            "into a robust network access-control platform for educational laboratories."
        ),
    ]

    elements += [
        heading(1, "References"),
        normal("[1] L. L. Peterson and B. S. Davie, Computer Networks: A Systems Approach, 6th ed. Morgan Kaufmann, 2021."),
        normal("[2] M. Grinberg, Flask Web Development: Developing Web Applications with Python, 2nd ed. O'Reilly Media, 2018."),
        normal("[3] K. Chodorow, MongoDB: The Definitive Guide, 3rd ed. O'Reilly Media, 2019."),
        normal("[4] M. Jones, J. Bradley, and N. Sakimura, JSON Web Token (JWT), RFC 7519, IETF, 2015."),
        normal("[5] D. F. Ferraiolo and D. R. Kuhn, \"Role-Based Access Controls,\" in Proc. 15th National Computer Security Conference, 1992, pp. 554-563."),
        normal("[6] P. Biondi and the Scapy community, Scapy Documentation, https://scapy.readthedocs.io/."),
        normal("[7] P. Mockapetris, Domain Names - Implementation and Specification, RFC 1035, IETF, 1987."),
        normal("[8] E. Rescorla, The Transport Layer Security (TLS) Protocol Version 1.3, RFC 8446, IETF, 2018."),
        normal("[9] D. Eastlake, Transport Layer Security (TLS) Extensions: Extension Definitions, RFC 6066, IETF, 2011."),
        normal("[10] R. Fielding and J. Reschke, Hypertext Transfer Protocol (HTTP/1.1): Message Syntax and Routing, RFC 7230, IETF, 2014."),
        normal("[11] N. Provos and D. Mazieres, \"A Future-Adaptable Password Scheme,\" in Proc. USENIX Annual Technical Conference, 1999, pp. 81-91."),
        normal("[12] Microsoft, \"netsh advfirewall firewall context,\" Microsoft Learn, https://learn.microsoft.com/windows-server/administration/windows-commands/netsh-advfirewall."),
        normal("[13] Pallets Projects, Flask Documentation, https://flask.palletsprojects.com/."),
        normal("[14] MongoDB Inc., MongoDB Manual, https://www.mongodb.com/docs/manual/."),
        heading(1, "Appendix A. Writing Compliance Checklist"),
        normal(
            "This appendix records how the draft follows the report-writing guide. The report uses the six-chapter structure required by the template, "
            "keeps the abstract as a single paragraph, includes acknowledgements within the requested length, introduces overview and conclusion paragraphs "
            "for technical chapters, cross-references figures and tables in body text, and uses IEEE-style numbered references. The final school template may "
            "already contain mandatory Appendix A content; if so, this checklist can be removed after copying the thesis content into the official editable file."
        ),
        table(
            "A.1",
            "Checklist for final editing before submission.",
            ["Item", "Status in this draft"],
            [
                ["Cover metadata", "Student name and ID are filled; supervisor, email, and program remain placeholders."],
                ["Abstract length", "Generated between 200 and 350 words."],
                ["Acknowledgements length", "Generated between 100 and 150 words."],
                ["Source basis", "Based on current source-derived report and static analysis; old documents are not used as the primary technical source."],
                ["Safety", "No agent executable, server, Docker, packet capture, netsh, or firewall code is executed by the generation script."],
            ],
        ),
        heading(1, "Appendix B. Main Use Case Specifications"),
        normal(
            "This appendix expands the most important use cases so that the main chapters can remain readable. The use cases are written at a level suitable for "
            "the thesis report and can be converted into diagrams in the final template if required by the supervisor."
        ),
        table(
            "B.1",
            "Detailed use case specifications.",
            ["Use case", "Precondition", "Main flow", "Postcondition"],
            [
                ["Dashboard login", "A user account exists and is active.", "User submits credentials; server validates password hash; session/JWT context is created; user reaches permitted pages.", "Authenticated user context is available and route guards apply RBAC."],
                ["Create whitelist entry", "User is authenticated and has access to the target scope.", "User submits domain/IP data; server validates input; service stores entry; version metadata is updated.", "Future agent synchronization can include the new entry."],
                ["Activate whitelist profile", "Teacher or admin has permission for the group.", "User selects a profile; server activates it for the group; effective whitelist state changes.", "Agents in the group receive updated profile information during synchronization."],
                ["Agent heartbeat", "Agent has valid identity and token.", "Agent sends status and metadata; server updates last-seen state; server may signal sync or policy changes.", "Dashboard can display current endpoint status."],
                ["Upload logs", "Agent has captured network metadata and holds a valid token.", "Agent sends a batch of logs; server stores documents in MongoDB; dashboard queries remain RBAC-filtered.", "Teachers and admins can review scoped network activity."],
            ],
        ),
    ]

    return elements


def build_long_elements() -> list[dict]:
    elements: list[dict] = []

    elements += [
        heading(1, "CHAPTER 1. INTRODUCTION"),
        heading(2, "1.1 Motivation"),
    ]
    add_paragraphs(elements, [
        "Computer laboratories in schools and universities are shared learning environments where the same machines are used by many classes, subjects, and teaching styles. In a single day, one room may support introductory programming, database practice, operating-system exercises, networking labs, online examinations, or general digital-literacy activities. These activities depend on Internet access, but their acceptable access boundaries are different. A programming class may need documentation websites, repositories, online judges, and package mirrors. An examination may need only the examination portal. A networking lesson may need controlled access to selected remote services. This changing context makes network access management a concrete educational problem rather than a purely technical preference.",
        "The problem becomes more serious because teachers usually do not have direct control over the network configuration of laboratory computers. If unrestricted access is available, students may open social networks, games, streaming services, or unrelated websites during class. If access is blocked too broadly, students may be unable to reach legitimate learning resources. Both situations reduce teaching effectiveness. The first situation distracts from the learning objective, while the second situation interrupts the lesson and forces teachers to wait for IT staff. A useful laboratory system must therefore balance access control and teaching flexibility.",
        "Existing operational practices often depend on manual configuration. An administrator may configure router rules, gateway filters, proxy settings, or endpoint firewall rules before a lesson. This practice is fragile because lesson requirements change frequently and because a small mistake can block necessary resources. It is also difficult to maintain a consistent policy across many machines. When a laboratory contains dozens of endpoints, updating every machine manually is time-consuming and error-prone. The situation is even more difficult when different teachers use different resources for different groups.",
        "The problem is not limited to blocking unwanted traffic. A complete educational-lab workflow also requires visibility and accountability. Teachers and administrators need to know whether endpoint agents are online, which machines belong to which group, what whitelist policy is active, and what network activity was observed during a session. Audit information is needed when administrative settings change. Without logs and role-aware visibility, network control becomes a hidden system that teachers cannot trust or explain.",
        "For these reasons, the thesis focuses on a practical problem: how to manage network access in educational computer laboratories in a way that is centralized for administrators, usable for teachers, enforceable on Windows endpoints, and observable through logs and audit records. The motivation of the thesis is the need for a system that supports classroom-specific network control without requiring every teacher to become a firewall administrator."
    ])

    elements += [heading(2, "1.2 Objectives and scope of the graduation thesis")]
    add_paragraphs(elements, [
        "Before defining the objective of the graduation thesis, it is necessary to compare the target problem with common existing approaches. Classroom-management products can provide screen monitoring, remote control, or activity supervision, but many of them are commercial, vendor-dependent, and difficult to adapt to a custom laboratory workflow. DNS-filtering products can block domain categories, but they often operate at the network or resolver level and do not naturally represent teacher-group ownership. Proxy-based solutions can provide centralized filtering and logging, but they require client configuration and do not cover every traffic type. Manual endpoint firewall rules use built-in operating-system mechanisms, but they are not manageable at classroom scale.",
        "The current project is positioned between these approaches. It does not attempt to become a full classroom-management suite or a commercial secure web gateway. Instead, it focuses on a specific set of needs: central management of agents, groups, users, API keys, whitelists, logs, audit events, and lesson-oriented profiles; role-based access control for administrators and teachers; and local endpoint enforcement on Windows machines. This focus is suitable for a graduation thesis because it requires distributed-system design, backend API implementation, endpoint programming, security reasoning, and practical deployment considerations.",
    ])
    elements.append(table(
        "1.1",
        "Comparison between SAINT and common network-control approaches.",
        ["Criterion", "Manual firewall", "DNS filter", "Proxy filter", "Classroom suite", "SAINT prototype"],
        [
            ["Teacher group workflow", "Weak", "Weak", "Medium", "Medium", "Strong through group and teacher RBAC"],
            ["Endpoint-specific policy", "Possible but manual", "Limited", "Limited", "Vendor-specific", "Designed through registered agents"],
            ["Central dashboard", "Usually absent", "Available in some products", "Available in some products", "Available", "Implemented through Flask dashboard"],
            ["Whitelist-only enforcement", "Possible but risky", "DNS-only", "Protocol-dependent", "Vendor-specific", "Implemented on Windows endpoint"],
            ["Audit and logs", "Weak", "Product-dependent", "Medium", "Product-dependent", "Implemented through audit and log models"],
            ["Customizability", "High but low-level", "Limited", "Medium", "Low for closed products", "High because source is project-owned"],
        ],
    ))
    add_paragraphs(elements, [
        "The first objective is to design a server that can act as the authority for network policy and operational data. The server must authenticate dashboard users, authorize actions according to roles, store persistent state in a database, provide REST APIs, render dashboard pages, and notify the frontend when relevant data changes. It must also support endpoint agents that are not human users and therefore need an enrollment mechanism based on API keys and tokens.",
        "The second objective is to design a Windows endpoint agent that can participate in this central-management model. The agent must register with the server, keep its identity and token state, send heartbeat information, synchronize whitelist data, resolve domains into addresses, update local firewall rules, capture useful network metadata, and upload logs. The agent must also provide a desktop interface so that its local status can be inspected without reading log files manually.",
        "The third objective is to support educational authorization. The system must distinguish an administrator from a teacher. An administrator should manage global system state, while a teacher should work only with assigned groups. This objective is important because laboratory systems often serve multiple classes, and a teacher should not accidentally inspect or modify another class's data. The thesis therefore treats RBAC and group filtering as core requirements rather than as cosmetic dashboard restrictions.",
        "The scope of the thesis is a source-level prototype. The system targets Windows endpoint agents, a Python Flask server, MongoDB persistence, and web dashboard access. The enforcement model is whitelist-based, meaning that the system is oriented toward allowing known educational resources rather than classifying all unsafe traffic. The thesis does not include HTTPS decryption, cross-platform endpoint support, commercial-grade update signing, enterprise identity federation, or large-scale production monitoring. These items are reserved for future work."
    ])

    elements += [heading(2, "1.3 Tentative solution")]
    add_paragraphs(elements, [
        "The selected direction is a distributed client-server architecture. The server is responsible for policy authority, authentication, authorization, persistence, dashboard operations, and auditability. The endpoint agent is responsible for applying policy locally and reporting runtime state. This direction was selected because the classroom problem requires both centralized management and local enforcement. A purely central gateway does not naturally express per-machine state, while a purely local tool is difficult to manage for many machines.",
        "The server side is implemented with Flask, Flask-SocketIO, PyMongo, Pydantic, JWT-related libraries, and bcrypt. The source code follows a controller-service-model style. Controller modules expose HTTP routes, service modules implement business logic, and model modules wrap MongoDB collections and indexes. MongoDB stores users, sessions, API keys, groups, agents, policies, whitelists, profiles, logs, audit records, and revoked tokens. SocketIO is used as a supporting real-time mechanism for the dashboard.",
        "The agent side is implemented in Python for Windows. It uses PySide6 for the desktop interface, requests for server communication, cryptography for configuration protection, DNS-related libraries for name resolution, Scapy for packet capture, pywin32 and firewall-related modules for Windows integration, and a controller/signal architecture to keep the GUI responsive. The agent does not expose a local HTTP API. It is a client of the central server.",
        "The main contribution of the tentative solution is the integration of five technical ideas into a coherent educational-lab prototype: whitelist-based endpoint firewall enforcement, versioned whitelist synchronization, teacher-scoped RBAC, lesson-oriented whitelist profiles, and packet metadata collection with domain extraction. The later chapters describe the requirements, technologies, detailed design, implementation, and contribution analysis of these ideas."
    ])

    elements += [heading(2, "1.4 Thesis organization")]
    add_paragraphs(elements, [
        "The remainder of the thesis is organized according to the graduation-thesis template. Chapter 2 presents requirement survey and analysis, including status survey, functional overview, use case diagrams, business process, functional descriptions, and non-functional requirements. Chapter 3 summarizes the theoretical background and technologies used by the project, explaining why each technology is relevant to the requirements. Chapter 4 presents architecture design, detailed design, application building, testing, and deployment. Chapter 5 discusses the main solution contributions in more depth, with each contribution analyzed in terms of problem, solution, and result. Chapter 6 concludes the thesis and describes future work. The references and appendices provide source citations, writing-template compliance notes, use case details, API catalogues, and module inventory."
    ])

    elements += [
        heading(1, "CHAPTER 2. REQUIREMENT SURVEY AND ANALYSIS"),
    ]
    add_paragraphs(elements, [
        "Chapter 1 identified the need for a classroom-oriented network access management system. This chapter refines that need into concrete requirements. The chapter first surveys the current situation and compares typical existing approaches. It then summarizes the actors and high-level functions of the system, presents use case diagrams and a business process, specifies important use cases, and defines non-functional requirements. The goal is to make the later design decisions traceable to educational and operational needs."
    ])
    elements += [heading(2, "2.1 Status survey")]
    add_paragraphs(elements, [
        "The status survey is based on three sources of requirements: classroom operation needs, existing categories of network-control products, and the source code of the current SAINT prototype. In the classroom context, the most important users are administrators and teachers. Administrators need maintainable system-wide control. Teachers need a limited but practical ability to prepare and activate policies for their own classes. Endpoint machines need a reliable way to receive and enforce the selected policy. These needs are narrower than a complete enterprise-security platform but broader than a simple block list.",
        "A manual firewall approach is familiar because Windows Firewall is already present on the target operating system. However, direct manual configuration is not appropriate for frequent classroom changes. Firewall rules have low-level syntax, may require administrator permission, and can easily block important connectivity if applied incorrectly. A teacher should not need to know how to create inbound and outbound firewall rules, how to restore a firewall snapshot, or how to translate domains into IP ranges before a lesson.",
        "DNS filtering is attractive because domain names are easier to understand than IP addresses. A DNS-based system can block or allow domains at the resolver level. However, DNS filtering has limitations in a shared laboratory. It may apply to an entire network rather than to a teacher's assigned group, and it may not provide endpoint-level information such as agent identity, local state, or packet metadata. It also cannot fully represent policy for traffic that does not use the expected resolver.",
        "Proxy filtering can provide centralized policy and logs for web traffic. Nevertheless, a proxy solution requires client configuration, browser or operating-system settings, and careful handling of applications that do not use the proxy. It also tends to focus on HTTP or HTTPS flows, while a laboratory may need a broader endpoint view. For a graduation thesis prototype, implementing and operating a secure proxy would also shift the focus away from distributed endpoint management.",
        "Commercial classroom-management tools often include rich classroom features such as screen monitoring, file distribution, messaging, or remote control. These products can be useful, but they may be costly, closed-source, and difficult to integrate with a custom role model. They may not expose the exact backend and agent behavior needed for research and thesis evaluation. The SAINT project instead emphasizes an inspectable implementation where the server, database models, agent runtime, and security behavior can be analyzed from source code.",
        "From the project source, the implemented system already contains many features required by this status survey: Flask controllers under an /api prefix, server-side services and models, MongoDB collections, JWT service, API-key service, RBAC configuration, dashboard templates, SocketIO support, PySide6 agent GUI, token manager, whitelist synchronization, firewall manager, DNS resolver, packet capture, log sender, heartbeat sender, encrypted configuration support, server tests, Docker-related files, and a PyInstaller specification. The requirements in this chapter are therefore not speculative; they are aligned with a real source tree."
    ])
    elements.append(table(
        "2.1",
        "Survey of existing approaches and requirements extracted for SAINT.",
        ["Approach", "Observed advantage", "Observed limitation", "Requirement derived for SAINT"],
        [
            ["Manual endpoint firewall", "Uses built-in Windows capability", "Hard to update safely across many machines", "Provide centralized policy and agent-managed firewall updates"],
            ["DNS filtering", "Domain-level policy is easy to understand", "Weak endpoint identity and group-specific classroom workflow", "Bind effective policy to registered agents and groups"],
            ["Proxy filtering", "Central logs and web filtering", "Client configuration and protocol coverage issues", "Avoid requiring proxy configuration for the prototype"],
            ["Commercial classroom suite", "Rich teaching workflow", "Closed, costly, and not source-inspectable", "Implement project-owned modules and document source behavior"],
            ["Standalone packet monitor", "Useful visibility", "No enforcement or teacher workflow", "Combine monitoring with whitelist enforcement and RBAC"],
        ],
    ))
    add_paragraphs(elements, [
        "The survey leads to a requirement that the project should be centralized in management but distributed in enforcement. A teacher should not manage local firewall rules directly. An administrator should not need to visit every endpoint to update the policy. An endpoint should not decide its own classroom whitelist independently of the server. These constraints motivate a model where the server stores and authorizes policy while the agent applies the effective policy locally.",
        "The survey also shows the need for observability. A network-control system that only blocks traffic is insufficient for educational use. Administrators need to know whether agents are online, whether API keys are used, whether user actions were audited, and whether logs are arriving. Teachers need to inspect only their assigned groups. Therefore, the system must store structured data and provide role-filtered dashboard access."
    ])

    elements += [
        heading(2, "2.2 Functional Overview"),
    ]
    add_paragraphs(elements, [
        "The functional overview summarizes the system at a high level before the detailed use case descriptions. SAINT has three operational actors: Administrator, Teacher, and Windows Agent. The Administrator owns system-wide resources such as users, API keys, groups, global whitelist entries, agents, audit logs, and overall settings. The Teacher owns classroom operations for assigned groups, especially whitelist profiles and scoped log review. The Windows Agent is an endpoint client that registers, synchronizes, enforces, monitors, and reports.",
        "The system functions can be grouped into dashboard functions and agent functions. Dashboard functions include authentication, user management, group management, agent management, whitelist management, whitelist profile management, log browsing, statistics, export, API-key management, and audit review. Agent functions include registration, token management, heartbeat, effective whitelist synchronization, DNS resolution, firewall rule management, packet capture, domain extraction, configuration protection, and log upload.",
        "The functional overview is intentionally broad. Detailed event flows, preconditions, and postconditions are specified later in Section 2.3. The overview shows that SAINT is not a single-purpose firewall utility. It is a distributed management system whose functions interact through identity, group assignment, policy state, endpoint runtime state, and stored logs."
    ])
    elements.append(table(
        "2.2",
        "Actors and high-level responsibilities.",
        ["Actor", "High-level responsibility", "Source-oriented implementation area"],
        [
            ["Administrator", "System-wide management of users, API keys, groups, agents, whitelist entries, policies, logs, and audit records.", "server/controllers, server/services, server/models"],
            ["Teacher", "Management of assigned groups, lesson whitelist profiles, scoped agents, and scoped logs.", "RBAC service, group service, whitelist profile service"],
            ["Windows Agent", "Registration, heartbeat, synchronization, firewall enforcement, packet monitoring, and log sending.", "agent/core, agent/controllers, agent/whitelist, agent/firewall, agent/capture"],
        ],
    ))
    elements += [heading(3, "2.2.1 General use case diagram")]
    add_paragraphs(elements, [
        "Figure 2.1 presents the general use case diagram for the system. The diagram separates human dashboard users from the endpoint agent because they authenticate and behave differently. Administrators and teachers use the web dashboard. Agents use server APIs as machine clients. The diagram also shows that teacher operations are intentionally narrower than administrator operations.",
    ])
    elements.append(fig("2.1"))
    add_paragraphs(elements, [
        "The Administrator actor is associated with user management, API-key management, group management, global whitelist management, agent management, log review, and audit review. These functions require broad authority because they affect the whole system. The Teacher actor is associated with scoped dashboard data, assigned groups, whitelist profile operations, and log review for classroom sessions. The Windows Agent actor is associated with registration, heartbeat, whitelist synchronization, firewall enforcement, and log upload.",
        "This general diagram also clarifies that students are not direct dashboard users in the current prototype. Students use laboratory machines that are affected by agent enforcement. This design choice simplifies the role model. The system focuses on administrators, teachers, and endpoint agents instead of building a separate student portal."
    ])
    elements += [heading(3, "2.2.2 Detailed use case diagram")]
    add_paragraphs(elements, [
        "Figure 2.2 decomposes the whitelist and classroom-policy use case because this is the most important functional area of the thesis. Whitelist management is not only create-read-update-delete behavior. It includes validation, import/export, profile management, activation, synchronization, DNS resolution, and local enforcement. The teacher-facing workflow and the agent-facing workflow meet at the server's effective policy calculation.",
    ])
    elements.append(fig("2.2"))
    add_paragraphs(elements, [
        "The detailed use case starts when a teacher or administrator prepares entries. The server validates whether an entry is a domain, IP address, URL, or category-related record and stores it according to global or group scope. A profile can be created to represent a lesson or classroom scenario. When the profile is activated for a group, the effective policy for agents in that group changes. The next synchronization request can then return updated entries and metadata.",
        "The endpoint side of the detailed use case begins when an agent pulls policy. The server returns the effective whitelist, active profile information, and version metadata. The agent resolves domain entries and updates firewall rules. This makes the policy operational on the Windows machine. The decomposition is important because it separates human-friendly policy editing from machine-level enforcement."
    ])
    elements += [heading(3, "2.2.3 Business process")]
    add_paragraphs(elements, [
        "Figure 2.3 illustrates a typical business process for using the system in a classroom. The process combines multiple use cases and therefore differs from a single use case event flow. It begins before the lesson, when the teacher or administrator prepares a whitelist profile and ensures that agents are assigned to the correct group. The policy is then activated, agents synchronize, endpoint firewalls enforce the policy, and the teacher reviews logs after or during the session.",
    ])
    elements.append(fig("2.3"))
    add_paragraphs(elements, [
        "The process also includes supporting administrative actions. The administrator must maintain users, API keys, global rules, and group assignments. Audit records should be created for important management operations. These supporting actions are not part of every classroom session, but they are necessary for a trustworthy system.",
        "The business process shows why SAINT needs both server and agent modules. If only the server exists, the policy is visible but not enforced. If only the agent exists, local enforcement becomes difficult to manage across many machines. The process therefore motivates the distributed design that is described in Chapter 4."
    ])

    elements += [heading(2, "2.3 Functional description")]
    add_paragraphs(elements, [
        "This section specifies five important use cases. The selected use cases cover both human dashboard operations and machine-client operations. They are Dashboard authentication, Agent registration and heartbeat, Whitelist and profile management, Endpoint whitelist enforcement, and Network log collection with dashboard review. Together, these use cases represent the main value of the project and the main risks that must be controlled."
    ])
    elements.extend(long_use_case_tables())
    add_paragraphs(elements, [
        "The use case descriptions show that authentication and authorization are prerequisites for almost every important operation. Even agent operations require a trusted identity. Dashboard users are checked according to role and permission, while agents are checked according to API-key enrollment and JWT-based calls. This distinction avoids mixing human identity with machine identity.",
        "The descriptions also show that whitelist policy is a multi-step process. A dashboard user creates policy data, the server stores and versions the policy, the agent synchronizes the effective result, and the firewall manager applies local rules. A failure at any step must be handled carefully. The system must not silently assume that a dashboard edit has already been enforced on every endpoint.",
        "The monitoring use case is also important because logs provide feedback. Without logs, the system would only state that a policy is configured. With logs, teachers and administrators can evaluate what happened on the endpoint, whether agents are active, and whether the policy aligns with classroom expectations."
    ])

    elements += [heading(2, "2.4 Non-functional requirement")]
    add_paragraphs(elements, [
        "Security is the first non-functional requirement. The server must protect dashboard endpoints and agent endpoints using appropriate authentication. Dashboard users require password-based login, session or token state, role checks, and permission checks. Agents require API-key enrollment and token-protected communication after registration. Passwords and sensitive keys must not be stored in plaintext. Audit logging should record important administrative actions so that system changes are accountable.",
        "Authorization is a separate requirement from authentication. A user may be authenticated but still not authorized to operate on a particular group. The teacher role requires data filtering by group ownership. This filtering must be applied in services and controllers that return groups, agents, logs, whitelist entries, and profiles. If filtering is implemented only in the user interface, unauthorized data may still be reachable through APIs. The source code therefore treats RBAC as backend logic.",
        "Reliability is critical because the endpoint agent can change firewall behavior. A failed firewall update can affect network access. The system must maintain self-allow rules, allow access to the server, preserve DNS reachability, clean up project-owned rules, and restore previous state when necessary. The agent should also preserve useful state across restarts, including configuration, identity, token information, and cached policy data where appropriate.",
        "Usability is required for both teachers and endpoint operators. Teachers should interact with groups, profiles, and whitelist entries rather than low-level firewall syntax. The agent GUI should display status, logs, settings, whitelist state, and firewall information in a way that supports troubleshooting. The GUI must remain responsive even when background workers perform synchronization or packet capture.",
        "Performance requirements are moderate but still important. The server should handle dashboard queries and agent requests without unnecessary blocking. MongoDB indexes should support common lookup fields such as agent_id, group_id, username, role, timestamp, key hash, and token identifiers. The agent should avoid excessive CPU usage during packet capture and avoid repeated firewall updates when whitelist versions have not changed.",
        "Maintainability is addressed through modular design. The server separates controllers, services, models, middleware, configuration, views, and tests. The agent separates GUI views, controller logic, lifecycle management, token management, whitelist synchronization, firewall management, capture, logging, network utilities, and configuration. This separation makes the system easier to reason about and reduces the chance that a GUI change accidentally affects firewall behavior.",
        "Deployability is also a requirement. The server includes requirements, Dockerfile, docker-compose, and .env-example files. The agent includes a PyInstaller specification for producing a Windows executable. However, production deployment requires additional hardening such as TLS, strong secrets, database backups, service monitoring, signed updates, and a staged rollout plan.",
    ])
    elements.append(table(
        "2.9",
        "Summary of non-functional requirements.",
        ["Requirement", "Concrete expectation", "Design implication"],
        [
            ["Security", "Protect dashboard and agent endpoints", "JWT, API keys, password hashing, token revocation, and audit logs"],
            ["Authorization", "Limit teacher access to assigned groups", "Backend RBAC and group ownership filtering"],
            ["Reliability", "Avoid accidental network loss", "Self-allow, server allow, DNS allow, snapshot, restore, and cleanup"],
            ["Usability", "Hide firewall complexity from teachers", "Dashboard profiles and agent GUI status views"],
            ["Performance", "Avoid unnecessary sync and slow queries", "Versioned sync and MongoDB indexes"],
            ["Maintainability", "Keep responsibilities separated", "Controller-service-model server and signal-driven agent"],
            ["Deployability", "Support reproducible setup", "Docker-related server files and PyInstaller agent spec"],
        ],
    ))
    add_paragraphs(elements, [
        "This chapter has transformed the general motivation into concrete requirements. The analysis shows that SAINT must include central management, endpoint enforcement, RBAC, policy versioning, monitoring, auditability, and deployment support. Chapter 3 next presents the theoretical background and technologies used to satisfy these requirements."
    ])

    elements += [heading(1, "CHAPTER 3. THEORETICAL BACKGROUND AND TECHNOLOGIES")]
    add_paragraphs(elements, [
        "Chapter 2 defined the requirements for the system. This chapter presents the theoretical background and technologies selected to satisfy those requirements. The chapter is organized by architectural concern rather than by source folder. For each technology, the discussion explains which requirement it supports and why it is suitable for the current prototype."
    ])
    elements += [heading(2, "3.1 Distributed client-server architecture")]
    add_paragraphs(elements, [
        "A distributed client-server architecture separates the authority that stores and decides policy from the clients that consume or enforce that policy. In SAINT, the server stores users, groups, agents, whitelist entries, profiles, logs, sessions, audit records, and API keys. Browser clients use the server for dashboard functions, while endpoint agents use the server for registration, heartbeat, synchronization, and log upload. This architecture directly supports the requirement that administrators manage policy centrally while endpoint machines enforce policy locally.",
        "Alternative approaches include a purely local agent, a gateway-only appliance, or a peer-to-peer design. A purely local agent would be difficult to manage at laboratory scale. A gateway-only design would not provide endpoint-local state or a direct Windows firewall integration. A peer-to-peer design would add complexity without matching the classroom workflow. Therefore, the client-server model is the most appropriate architectural baseline for the thesis."
    ])
    elements += [heading(2, "3.2 Flask, REST APIs, and server-rendered dashboard")]
    add_paragraphs(elements, [
        "Flask is used as the web framework for the server. It is lightweight, familiar in Python projects, and flexible enough to support an application factory, blueprints, middleware, templates, and extensions. REST APIs are used because most operations are resource-oriented: create a user, list agents, update a group, register an agent, send heartbeat, synchronize whitelist entries, receive logs, or revoke an API key. The REST style makes request boundaries explicit and supports authentication decorators on routes.",
        "The project also uses server-rendered dashboard templates. This choice is pragmatic for a graduation-thesis prototype because it avoids the need for a separate frontend build system while still providing web pages for agents, groups, whitelist, logs, API keys, users, audit, login, profile, and password change. SocketIO can then be used as an additional real-time mechanism without replacing the basic HTTP workflow.",
        "A possible alternative would be a separate single-page application with a JavaScript framework such as React or Vue. That approach could provide a richer frontend but would increase development scope. Another alternative would be a pure API server with no dashboard templates, but that would make demonstration and classroom operation harder. The selected design gives a complete but manageable web interface."
    ])
    elements += [heading(2, "3.3 MongoDB and document modeling")]
    add_paragraphs(elements, [
        "MongoDB is selected as the persistence layer because many entities in SAINT are naturally document-shaped. Agents can contain host metadata, status, heartbeat timestamps, group identifiers, and runtime properties. Logs can contain packet metadata, domains, addresses, timestamps, and classification fields. Whitelist entries can represent domains, IP addresses, URLs, categories, active status, ownership scope, and version-related metadata. A document database can store these records without forcing every entity into a rigid relational schema.",
        "The source code uses model classes to wrap collection access and index creation. This design keeps database operations out of controllers and allows services to express business rules at a higher level. The identified collections are agents, agent_policies, api_keys, audit_logs, groups, logs, admin_sessions, users, whitelist, whitelist_meta, whitelist_profiles, and revoked_tokens. Indexes are important for fields that are frequently queried, such as agent identifiers, device identifiers, group identifiers, usernames, roles, timestamps, API-key hashes, and token identifiers.",
        "A relational database could also be used for this project, especially for users, groups, and permissions. However, log documents and agent metadata may evolve during development. MongoDB makes it easier to store these records during prototyping. The trade-off is that relationships and constraints must be enforced carefully in application services rather than by foreign keys alone."
    ])
    elements += [heading(2, "3.4 JWT, API keys, password hashing, and RBAC")]
    add_paragraphs(elements, [
        "SAINT uses different authentication mechanisms for different actors. Human dashboard users authenticate with credentials and receive an authenticated context. Endpoint agents register with API keys and later call protected APIs with JWT tokens. JWT is suitable for representing identity and claims in a compact form [4]. API keys are suitable for controlled enrollment because an administrator can issue, validate, revoke, and scope them. Password hashing through bcrypt-style mechanisms protects stored user passwords [11].",
        "RBAC is used to separate administrators and teachers. Classical role-based access control assigns permissions to roles and users to roles [5]. SAINT extends this practical role separation with group ownership filtering. A teacher can be authenticated and still be denied access to a group that is not assigned to that teacher. This distinction is necessary because classroom systems involve multiple groups and potentially many teachers.",
        "Alternative authorization models include discretionary access control, attribute-based access control, or hard-coded route checks. Discretionary access control is not suitable because classroom ownership should not depend on arbitrary user-granted permissions. Attribute-based access control is powerful but may be too complex for this prototype. RBAC with group filtering is a practical middle ground that matches the current role set."
    ])
    elements += [heading(2, "3.5 Windows Firewall, DNS, and whitelist enforcement")]
    add_paragraphs(elements, [
        "Windows Firewall is used as the local enforcement mechanism because the target endpoints are Windows laboratory machines and the firewall is already integrated into the operating system. The agent can create allow rules, manage project-owned rules, and support restrictive behavior. This choice satisfies the requirement that enforcement happen on the endpoint rather than only at a central dashboard.",
        "Domain-based policy requires DNS resolution because firewall rules typically operate on IP addresses or program/network conditions. DNS is the standard naming system for resolving domain names [7]. However, DNS resolution can return multiple addresses, and addresses can change over time. The agent therefore needs a resolver component and a synchronization process that can refresh the effective policy when necessary.",
        "Whitelist-only enforcement is safer for classroom focus than a blacklist-only approach because the teacher usually knows the resources needed for a lesson. A blacklist must predict every unwanted site, while a whitelist defines the allowed learning scope. The trade-off is operational risk. If a required service is omitted from the whitelist, the lesson may be interrupted. The agent must therefore include self-allow, server allow, DNS allow, snapshot, restore, and cleanup behavior."
    ])
    elements += [heading(2, "3.6 Packet capture and domain extraction")]
    add_paragraphs(elements, [
        "Packet capture provides visibility into network activity. The agent uses Scapy-related functionality to inspect packet metadata [6]. The purpose is not to decrypt content but to collect useful information such as addresses, protocols, ports, timestamps, and domains when available. Domain information can come from DNS queries, HTTP Host headers, and TLS Server Name Indication. HTTP/1.1 defines the Host header [10], while TLS extensions define SNI [9].",
        "This design supports monitoring without introducing full content inspection. It also matches the educational context, where the goal is to know whether students are reaching relevant or irrelevant domains rather than to inspect private content. The limitation is that some encrypted or privacy-preserving traffic may hide domain information. The log analysis must therefore be interpreted as metadata-based observation, not complete content visibility."
    ])
    elements += [heading(2, "3.7 PySide6 and signal-driven desktop design")]
    add_paragraphs(elements, [
        "The agent GUI uses PySide6, the Qt binding for Python. A desktop GUI is useful because the endpoint operator may need to inspect registration state, server connectivity, firewall status, whitelist entries, logs, or settings. The source architecture separates GUI views from background worker logic through an agent controller and signal bridge. Worker threads emit state changes, and the GUI thread updates widgets through queued signals.",
        "A signal-driven design is important for responsiveness. Packet capture, DNS resolution, heartbeat, log upload, and firewall updates should not block the interface. The agent code includes a queue and bridge mechanism so that updates can be drained on the Qt main thread. This approach is closer to an MVP-style design than a monolithic script where GUI code directly performs all network and firewall work."
    ])
    elements.append(table(
        "3.1",
        "Technology selection mapped to requirements.",
        ["Requirement", "Selected technology", "Reason for selection", "Alternative considered"],
        [
            ["Central API and dashboard", "Flask and server-rendered templates", "Lightweight Python web framework and manageable thesis scope", "Django or separate SPA frontend"],
            ["Real-time dashboard updates", "Flask-SocketIO with gevent dependencies", "Complements REST for event-style updates", "Polling-only dashboard"],
            ["Document persistence", "MongoDB with PyMongo", "Flexible documents for agents, logs, whitelist, and sessions", "Relational database"],
            ["Human and agent authentication", "JWT, API keys, bcrypt", "Separates dashboard users from machine enrollment", "Single static shared secret"],
            ["Authorization", "RBAC with group filtering", "Matches administrator and teacher responsibilities", "Hard-coded UI-only checks"],
            ["Endpoint enforcement", "Windows Firewall integration", "Uses built-in Windows capability", "Gateway-only firewall"],
            ["Monitoring", "Scapy packet capture", "Captures metadata on endpoint", "Passive server-only logs"],
            ["Agent GUI", "PySide6", "Native desktop interface with signal support", "Command-line-only agent"],
        ],
    ))
    add_paragraphs(elements, [
        "This chapter has introduced the technologies used in the project and explained how each choice supports the requirements from Chapter 2. The next chapter applies these technologies to the concrete architecture, design, implementation, testing, and deployment of SAINT."
    ])

    elements += [heading(1, "CHAPTER 4. DESIGN, IMPLEMENTATION, AND EVALUATION")]
    add_paragraphs(elements, [
        "Chapter 3 explained the technologies used by the project. This chapter presents how those technologies are assembled into the implemented system. The chapter follows the graduation-thesis template: architecture design, detailed design, application building, testing, and deployment. The descriptions are based on static source analysis and the source-generated documentation. No agent executable, GUI runtime, server runtime, Docker service, packet capture, netsh command, or firewall-changing code was executed while preparing this report."
    ])
    elements += [heading(2, "4.1 Architecture design"), heading(3, "4.1.1 Software architecture selection")]
    add_paragraphs(elements, [
        "The server uses a controller-service-model organization that is similar to MVC in spirit but adapted to a Flask API and dashboard project. Controllers define blueprints, parse requests, apply authentication or authorization decorators, and return responses. Services contain business logic such as registering agents, validating API keys, filtering teacher data, updating whitelist profiles, storing logs, and writing audit records. Models encapsulate MongoDB collections, indexes, and persistence operations. Templates and static files form the presentation layer for the web dashboard.",
        "The agent uses a model-view-presenter-like organization. The PySide6 views display state and collect user actions. The AgentController coordinates lifecycle and worker components. The signal bridge transfers state changes from background workers to the Qt main thread. Runtime components such as the token manager, whitelist manager, firewall manager, packet sniffer, heartbeat sender, log sender, and configuration manager operate outside the view layer. This design was selected because it reduces coupling between the GUI and risky system operations.",
        "A microservice architecture was not selected because the project is a graduation-thesis prototype and the current domain can be implemented coherently in one server process with modular packages. Splitting authentication, agent management, logs, and whitelist into separate services would increase deployment complexity without providing immediate educational value. A monolithic script was also rejected because it would make testing, source analysis, and future maintenance difficult."
    ])
    elements += [heading(3, "4.1.2 Overall design")]
    add_paragraphs(elements, [
        "Figure 4.1 shows the package-level architecture. The top layer contains presentation components: server-rendered web templates for dashboard users and the PySide6 GUI for the Windows agent. The controller layer contains Flask blueprints on the server and the AgentController on the endpoint. The service layer contains authentication, RBAC, whitelist, log, group, API-key, audit, and agent services. The persistence layer is represented by MongoDB models and collections. Endpoint runtime packages provide firewall, DNS, capture, token, configuration, and log-sending functions.",
    ])
    elements.append(fig("4.1"))
    add_paragraphs(elements, [
        "The dependency direction is intentionally controlled. Server controllers depend on services, and services depend on models. Models do not depend on controllers. Dashboard templates do not directly access MongoDB. On the agent side, views depend on controller signals rather than directly invoking firewall operations. Runtime workers communicate status through the controller and signal bridge. This dependency structure makes it easier to identify where a security decision or side effect occurs.",
        "The overall design also separates human and machine clients. The browser client is operated by an administrator or teacher and uses dashboard routes and APIs. The agent client is operated by endpoint software and uses registration, heartbeat, synchronization, and log-upload APIs. The server must handle both clients but must not confuse their authentication models."
    ])
    elements += [heading(3, "4.1.3 Detailed package design")]
    add_paragraphs(elements, [
        "The Windows agent package design is especially important because it contains the components that can affect network access. Figure 4.2 shows the main package relationships. The GUI package contains views and reusable components. The controllers package contains the AgentController. The core package contains lifecycle, registry, handlers, and token-related logic. The whitelist package manages synchronization and state. The firewall package manages rules. The capture and logging packages observe network traffic and send records to the server.",
    ])
    elements.append(fig("4.2"))
    add_paragraphs(elements, [
        "The design isolates high-risk operations behind dedicated packages. Firewall changes are not scattered across GUI widgets. Whitelist synchronization is not mixed directly with packet capture. Token and configuration responsibilities are separated from the dashboard views. This separation is necessary because the agent must remain understandable and recoverable when a policy update fails or when connectivity is interrupted.",
        "The server package design follows a similar separation. Authentication controllers call authentication services and JWT services. Agent controllers call agent services and policy services. Whitelist controllers call whitelist services and RBAC services. Log controllers call log services and RBAC services. The service layer is therefore the main location where business rules should be enforced."
    ])
    elements += [heading(2, "4.2 Detailed design"), heading(3, "4.2.1 User interface design")]
    add_paragraphs(elements, [
        "The dashboard user interface is designed for administrators and teachers who need to scan operational information quickly. The important screens include login, dashboard overview, agents, groups, whitelist, logs, API keys, user administration, audit, profile, and password change. The dashboard should use consistent table layouts for list data, clear action buttons for create-update-delete operations, and role-aware visibility so that teachers see only actions relevant to assigned groups.",
        "The agent user interface is a Windows desktop interface implemented with PySide6. Its important views include dashboard status, firewall information, whitelist information, logs, and settings. The design goal is not a marketing-style interface but an operational tool. The user should see whether the agent is registered, whether it can reach the server, whether whitelist synchronization is current, whether firewall mode is active, and whether logs are being captured or sent.",
        "Feedback placement is important. Network and firewall actions may fail because of permissions, invalid configuration, unavailable server, DNS failure, or local system restrictions. The UI should therefore present status and errors near the related function. A status-only design would be insufficient because a user might see that the agent is offline without knowing whether the cause is token expiration, server URL error, or network interruption.",
        "The interface design also needs to respect background work. Packet capture and heartbeat run continuously, while the GUI must remain responsive. The signal bridge and queue mechanism support this requirement by draining worker updates on the GUI thread. This design avoids direct widget updates from worker threads and reduces unnecessary rendering through diff-skip behavior described in the source-based report."
    ])
    elements += [heading(3, "4.2.2 Layer design")]
    add_paragraphs(elements, [
        "The first important class group is the server authentication and JWT layer. Its responsibility is to authenticate dashboard users, issue or refresh tokens, revoke tokens, track sessions, and validate current-user context. Inputs include credentials, refresh tokens, access tokens, session identifiers, and user identifiers. Outputs include authenticated identity, role information, token metadata, and error responses. Side effects include session creation, token revocation, and audit records.",
        "The second important class group is the RBAC and group filtering layer. Its responsibility is to decide what a current user may access. Inputs include user role, user identifier, group identifiers, permission names, and requested resources. Outputs include authorization decisions or filtered query scopes. The key side effect is not database mutation but data visibility control. This layer is essential for teacher accounts because it prevents accidental access to unrelated classroom data.",
        "The third important class group is the whitelist and profile layer. Its responsibility is to store entries, validate changes, maintain version metadata, compute effective policy, and support profile activation. Inputs include domain, IP, URL, category, group identifier, profile identifier, active state, and version fields. Outputs include stored entries, effective whitelist lists, active profile data, and synchronization responses. Side effects include version changes that trigger agent updates.",
        "The fourth important class group is the agent firewall and synchronization layer. Its responsibility is to ask the server for policy, compare versions, resolve domains, update local firewall rules, and preserve recovery paths. Inputs include server URL, agent identity, token, known versions, whitelist entries, DNS results, and policy mode. Outputs include local rule state, synchronization status, and UI signals. Side effects include Windows Firewall changes, which is why this layer must be isolated and carefully tested.",
        "The fifth important class group is the capture and logging layer. Its responsibility is to capture network metadata, extract domains, buffer log records, and send batches to the server. Inputs include packet metadata, DNS payloads, TLS handshakes, HTTP headers, timestamps, local endpoint state, and authentication headers. Outputs include structured log records and upload responses. Side effects include database log insertion on the server."
    ])
    elements.append(table(
        "4.1",
        "Layer design summary for important components.",
        ["Layer or class group", "Main inputs", "Main outputs", "Important side effects"],
        [
            ["Authentication and JWT", "Credentials, tokens, sessions", "User identity, token metadata, auth errors", "Session creation and token revocation"],
            ["RBAC and group filtering", "Role, user ID, group IDs, permission name", "Authorization decision or filtered scope", "Restricts visibility for teacher users"],
            ["Whitelist and profile service", "Entries, profile state, group ID, version", "Stored policy and effective sync result", "Updates version metadata"],
            ["Agent sync and firewall", "Token, policy mode, entries, DNS results", "Local rule state and status signals", "Changes Windows Firewall rules"],
            ["Capture and logging", "Packet metadata and extracted domains", "Structured log batches", "Inserts logs into MongoDB"],
        ],
    ))
    add_paragraphs(elements, [
        "For communication between these layers, the most important sequence is agent enrollment and synchronization. Figure 4.4 shows that the agent first registers with an API key, the server stores or updates the agent record, and the server returns identity and token data. Later, the agent sends heartbeat and version information. The server can then return synchronization data or policy changes. Finally, the agent uploads captured log batches, and the server persists them in MongoDB.",
    ])
    elements.append(fig("4.4"))
    elements += [heading(3, "4.2.3 Database design")]
    add_paragraphs(elements, [
        "The database design uses MongoDB collections rather than relational tables. Figure 4.3 summarizes the logical relationship among important collections. Users and groups represent dashboard ownership. Agents belong to groups and produce logs. Whitelist entries and whitelist profiles belong to global or group scopes. API keys enroll agents. Sessions and revoked tokens manage authentication state. Audit logs record administrative actions. Agent policies store per-agent policy overrides or runtime state.",
    ])
    elements.append(fig("4.3"))
    elements.append(table(
        "4.2",
        "MongoDB collection responsibilities.",
        ["Collection", "Responsibility", "Important relationships"],
        [
            ["users", "Stores administrator and teacher accounts, role, password hash, and lock state.", "Related to groups through teacher ownership and sessions through login state."],
            ["groups", "Stores classroom or laboratory groups and teacher assignments.", "Related to agents, whitelist entries, profiles, and teacher visibility."],
            ["agents", "Stores endpoint identity, host metadata, group assignment, and heartbeat status.", "Related to groups, logs, and agent policies."],
            ["logs", "Stores network metadata uploaded by agents.", "Related to agents and filtered by group ownership."],
            ["whitelist", "Stores global or group-scoped allow entries.", "Related to groups and whitelist_meta versioning."],
            ["whitelist_profiles", "Stores lesson-oriented profile data.", "Related to groups, teachers, and active policy state."],
            ["api_keys", "Stores enrollment keys and key metadata.", "Used by agents during registration."],
            ["admin_sessions", "Stores dashboard session and token identifiers.", "Related to users and JWT lifecycle."],
            ["revoked_tokens", "Stores revoked token identifiers.", "Used by JWT validation logic."],
            ["audit_logs", "Stores management action history.", "Related to users and security accountability."],
            ["agent_policies", "Stores per-agent policy overrides or runtime policy state.", "Related to agents."],
            ["whitelist_meta", "Stores global whitelist version metadata.", "Used by synchronization logic."],
        ],
    ))
    add_paragraphs(elements, [
        "The database design must support both dashboard queries and agent operations. Dashboard queries often filter by group, role, timestamp, or search text. Agent operations often look up agent identity, token state, policy version, and group membership. Index setup in model classes is therefore an important part of the design. Without indexes, a laboratory with many logs could experience slow dashboard queries and delayed synchronization.",
        "The design also separates audit records from network logs. Network logs describe endpoint traffic metadata. Audit logs describe administrative operations. This separation is important because these records have different audiences, retention requirements, and query patterns. A teacher may need network logs for assigned groups, while an administrator may need audit logs to investigate who changed a policy."
    ])
    elements += [heading(2, "4.3 Application Building"), heading(3, "4.3.1 Libraries and Tools")]
    elements.append(table(
        "4.3",
        "Libraries and tools used to build SAINT.",
        ["Purpose", "Library or tool", "Version source", "Role in the project"],
        [
            ["Server framework", "Flask", "server/requirements.txt", "REST APIs and dashboard routes"],
            ["Realtime events", "Flask-SocketIO, gevent, gevent-websocket", "server/requirements.txt", "SocketIO and asynchronous event support"],
            ["Database", "PyMongo, MongoDB", "server/requirements.txt and model source", "Document persistence and indexes"],
            ["Validation and config", "Pydantic, python-dotenv", "server/requirements.txt", "Configuration and structured data support"],
            ["Security", "PyJWT, bcrypt", "server/requirements.txt", "Token handling and password hashing"],
            ["Agent GUI", "PySide6", "agent/requirements.txt", "Windows desktop user interface"],
            ["Agent HTTP", "requests, urllib3", "agent/requirements.txt", "Server communication"],
            ["DNS", "dnspython, aiodns", "agent/requirements.txt", "Domain resolution for whitelist entries"],
            ["Packet capture", "Scapy", "agent/requirements.txt", "Network metadata capture"],
            ["Windows integration", "pywin32, pydivert", "agent/requirements.txt", "Windows system and network integration"],
            ["System monitoring", "psutil, netifaces", "agent/requirements.txt", "Process and network interface information"],
            ["Configuration security", "cryptography", "agent/requirements.txt", "Encrypted configuration support"],
            ["Packaging", "Dockerfile, docker-compose, PyInstaller spec", "server and agent folders", "Deployment artifacts"],
        ],
    ))
    elements += [heading(3, "4.3.2 Achievement")]
    add_paragraphs(elements, [
        "The project achievement is a complete prototype with both server and endpoint source code. Static source analysis identified sixty-four Python modules in the agent folder and forty-nine Python modules in the server folder, including tests. The server exposes approximately seventy REST route declarations under the /api prefix. The database design includes twelve MongoDB collections. The server test folder includes tests for agents, whitelist and logs, authentication, teacher data filtering, groups, and audit behavior.",
        "The server achievement includes an application factory, controller blueprints, service classes, MongoDB models, RBAC configuration, JWT service, API-key service, dashboard templates, static frontend assets, Docker-related files, and environment example configuration. The dashboard routes support pages for core administrative workflows. The API routes support both dashboard clients and endpoint agents.",
        "The agent achievement includes a PySide6 GUI, central agent controller, signal bridge, lifecycle initialization, registry, handlers, token manager, whitelist synchronization, firewall manager, DNS resolver, packet sniffer, domain extraction, log sender, heartbeat sender, encrypted configuration support, network utilities, cache/state modules, and packaging specification. The agent is designed as a client of the server and does not expose a local HTTP API.",
        "A notable achievement is that the project connects several security-relevant capabilities instead of implementing only one isolated feature. Authentication, API-key enrollment, token management, RBAC, whitelist policy, firewall enforcement, packet monitoring, log storage, audit trail, and dashboard presentation are all represented in the source. This integration makes the project suitable as a graduation-thesis system."
    ])
    elements.append(table(
        "4.4",
        "Implementation statistics from static source analysis.",
        ["Metric", "Observed value", "Interpretation"],
        [
            ["Agent Python modules", "64", "Agent is split into GUI, controller, lifecycle, network, firewall, capture, logging, and configuration packages."],
            ["Server Python modules", "49 including tests", "Server includes application, controllers, services, models, middleware, config, and tests."],
            ["Server API routes", "Approximately 70", "REST API surface covers dashboard management and agent operations."],
            ["MongoDB collections", "12", "Persistence covers identity, policy, logs, sessions, audit, and token lifecycle."],
            ["Server test files", "7 groups identified", "Tests cover agents, auth, whitelist/logs, teacher filtering, groups, and audit."],
            ["Generated report diagrams", "9 figures in this report", "Figures are source-derived and static, not runtime screenshots."],
        ],
    ))
    elements += [heading(3, "4.3.3 Illustration of main functions")]
    add_paragraphs(elements, [
        "The first main function is agent management. The dashboard can list agents, inspect an agent, update group assignment, update display name, update position, retrieve policy, and set policy. The agent sends registration and heartbeat requests. This function connects the human view of a laboratory machine with the machine-client identity used by the endpoint agent.",
        "The second main function is whitelist management. Dashboard users can list, create, import, export, bulk update, bulk delete, and inspect whitelist statistics. Agents can call the synchronization endpoint to receive effective whitelist data. This function is central to the thesis because it connects teacher-friendly policy editing with endpoint enforcement.",
        "The third main function is whitelist profile management. Profiles allow lesson-specific policy sets for groups. The server includes routes to list profiles, create profiles, update profiles, delete profiles, activate profiles, deactivate profiles, and list a user's own profiles. This function supports classroom flexibility because a teacher can prepare a policy for a lesson without rewriting the base whitelist every time.",
        "The fourth main function is logging and statistics. Agents upload log batches, while dashboard users can list logs, clear logs, export logs, and view statistics. The log feature supports accountability and lesson review. It also provides feedback for future policy refinement because administrators can see whether important educational domains were blocked or whether irrelevant domains appeared during a session.",
        "The fifth main function is authentication and RBAC. Human users log in, update profile information, change password, refresh tokens, and log out. API keys manage agent enrollment. RBAC filters data for teachers. This function is essential because a system that controls network access must also control who can change policies."
    ])
    elements += [heading(2, "4.4 Testing")]
    add_paragraphs(elements, [
        "The source tree includes server-side tests that represent the most important backend behaviors. The tests were not executed while preparing this report because the report-generation task explicitly avoids running the server or any component that could interact with networking. The testing description is therefore based on static inspection of the test files. In a final evaluation environment, these tests should be executed with a controlled test database and isolated configuration.",
        "The first testing area is agent behavior. Tests for agent registration and heartbeat verify that the server can accept machine-client enrollment and update agent state. These tests are important because all endpoint functions depend on trusted agent identity. If registration fails or heartbeat state is wrong, the dashboard cannot reliably represent laboratory endpoints.",
        "The second testing area is whitelist and log behavior. These tests verify that whitelist operations and log receiving work through the server API. They are important because whitelist synchronization and log collection are the main server-agent workflows. A bug in these areas could result in stale policy or missing observation data.",
        "The third testing area is authentication and teacher data filtering. These tests verify that user authentication works and that teachers receive scoped data. This is one of the highest-risk backend features because a broken filter could expose another class's agents or logs. Testing RBAC behavior is therefore more important than testing visual dashboard layout.",
        "The fourth testing area is groups and audit. Group tests verify classroom organization and assignment behavior. Audit tests verify that management operations can be recorded. These functions support accountability and system maintenance."
    ])
    elements.append(table(
        "4.5",
        "Testing scope identified from source files.",
        ["Test area", "Representative behavior", "Reason it matters"],
        [
            ["Agent tests", "Registration, heartbeat, agent APIs", "Endpoint identity is the basis for synchronization and logs"],
            ["Whitelist and log tests", "Whitelist CRUD, sync, log receiving", "Policy and monitoring are central workflows"],
            ["Authentication tests", "User login and token behavior", "Dashboard operations require trusted identity"],
            ["Teacher filtering tests", "Scoped data access", "Teachers must not see unrelated groups"],
            ["Group tests", "Group CRUD and assignment", "Group ownership drives policy and RBAC"],
            ["Audit tests", "Audit record behavior", "Administrative changes need traceability"],
        ],
    ))
    elements += [heading(2, "4.5 Deployment")]
    add_paragraphs(elements, [
        "The server deployment artifacts include a Dockerfile, docker-compose file, requirements file, and .env-example. These files indicate that the server can be deployed in a container-oriented environment with external configuration. The report does not include secret values and does not read the real server .env file. A production deployment should define MongoDB connection settings, JWT secrets, CORS policy, TLS termination, log retention, and administrator account initialization securely.",
        "The agent deployment artifact is a PyInstaller specification for building a Windows executable. The packaged agent is intended to run on laboratory endpoints without requiring users to install Python manually. However, endpoint deployment requires careful operational preparation because the agent can modify firewall behavior. A staged rollout should begin on non-critical machines, with a recovery plan, local administrator access, and clear instructions for restoring network access.",
        "A safe deployment model would include three phases. In the first phase, the server is deployed with test data and the agent is run in a non-enforcing or dry-run mode. In the second phase, a small number of endpoints are assigned to a test group and whitelist synchronization is validated. In the third phase, whitelist-only enforcement is enabled only after required server, DNS, and learning-resource allow rules are confirmed. This staged process reduces the risk of accidentally interrupting a classroom network.",
        "The current prototype provides the source-level basis for deployment, but production hardening remains future work. Important additions include TLS everywhere, signed agent updates, better observability, backup procedures, policy simulation, automatic rollback, and a formal incident procedure for firewall misconfiguration."
    ])
    add_paragraphs(elements, [
        "This chapter has presented the architecture, detailed design, application-building results, testing scope, and deployment considerations of SAINT. It shows how the requirements from Chapter 2 and technologies from Chapter 3 are realized in a modular server and Windows endpoint agent. Chapter 5 next focuses on the main technical contributions and explains why they are important."
    ])

    elements += [heading(1, "CHAPTER 5. SOLUTION AND CONTRIBUTION")]
    add_paragraphs(elements, [
        "Chapter 4 described the full design and implementation. This chapter focuses on the technical contributions that are most important for evaluating the thesis. Each contribution is presented through three aspects: the problem that motivated it, the solution implemented in the source, and the result or value achieved by the prototype."
    ])
    contribution_sections = [
        ("5.1 Whitelist-based endpoint firewall enforcement",
         "A classroom whitelist is useful only if it can be enforced. If a teacher prepares a list of allowed resources but endpoint machines remain unrestricted, the whitelist is merely documentation. At the same time, enforcement is risky because a wrong firewall rule can interrupt the agent's own connectivity or block resources needed during class.",
         "The solution is to make the Windows agent responsible for translating effective server policy into local firewall behavior. The agent synchronizes whitelist entries, resolves domains, prepares allow rules, protects access to the server and DNS, and applies restrictive behavior when whitelist-only mode is enabled. The firewall-related logic is isolated in dedicated packages so that GUI code does not directly manipulate system rules.",
         "The result is a prototype that demonstrates endpoint-level policy enforcement controlled from a central server. The design is source-inspectable and includes risk-reduction mechanisms such as self-allow rules, server allow rules, snapshot/restore concepts, cleanup routines, and status reporting. This contribution is the security core of the project."),
        ("5.2 Versioned whitelist synchronization",
         "Endpoint agents should not repeatedly download and reapply policy when nothing has changed. Reapplying firewall rules unnecessarily can waste time, increase failure risk, and make troubleshooting difficult. In a classroom setting, policy changes should be timely but not noisy.",
         "The solution is version-aware synchronization. The server maintains whitelist and group-related version metadata, and the agent sends known version information when requesting policy. The server can return the effective whitelist, policy mode, active profile, and version data. The agent can then decide whether local updates are needed.",
         "The result is a cleaner synchronization model. It supports global policy, group policy, and active profile behavior while keeping endpoint updates controlled. This contribution also creates a foundation for future scheduling, simulation, and rollback features."),
        ("5.3 Teacher-scoped RBAC and group ownership filtering",
         "A multi-class laboratory system must prevent teachers from seeing or modifying unrelated classroom data. Basic login is not enough because an authenticated teacher may still request data outside the assigned group. UI-only restrictions are also insufficient because APIs could be called directly.",
         "The solution is backend RBAC combined with group ownership filtering. The server distinguishes administrator and teacher roles. Services and controllers apply current-user context and group assignment rules when returning groups, agents, logs, whitelist data, and profiles. Permissions are therefore enforced at the data layer rather than only in the visual interface.",
         "The result is a system that models realistic educational responsibility. Administrators retain global authority, while teachers operate within assigned groups. The source also includes tests for teacher data filtering, indicating that the project treats this as a behavioral requirement rather than a presentation detail."),
        ("5.4 Lesson-oriented whitelist profiles",
         "A static whitelist is not flexible enough for different lessons. A teacher may need one policy for a programming exercise, another for an online quiz, and another for a network experiment. Editing the base whitelist for every lesson is error-prone and makes it difficult to return to a previous configuration.",
         "The solution is the whitelist profile feature. Profiles can represent lesson-specific allow lists for groups. The server provides profile operations such as listing, creating, updating, deleting, activating, and deactivating. The active profile is included in effective policy information used by agents.",
         "The result is a classroom-oriented policy mechanism. Teachers can prepare and activate profiles without rewriting all whitelist data. This contribution connects technical policy management to real teaching workflow and prepares the system for future scheduled or exam-mode profiles."),
        ("5.5 Packet metadata capture and domain extraction",
         "Network enforcement should be accompanied by visibility. A teacher or administrator needs to know what traffic was observed, which domains appeared, and whether endpoint activity matches the lesson. Raw IP addresses alone are often difficult to interpret.",
         "The solution is endpoint packet capture with domain extraction where possible. The agent captures network metadata and extracts domain information from DNS, HTTP Host headers, and TLS SNI. The log sender batches records and uploads them to the server, where logs are stored in MongoDB and shown through scoped dashboard routes.",
         "The result is an observation layer that complements enforcement. The system does not decrypt application content, but it collects useful metadata for classroom supervision and policy refinement. This contribution helps administrators and teachers understand the effect of whitelist policies."),
    ]
    for title, problem, solution, result in contribution_sections:
        elements.append(heading(2, title))
        add_paragraphs(elements, [
            f"(i) Problem. {problem}",
            f"(ii) Solution. {solution}",
            f"(iii) Result. {result}",
        ])
    elements.append(fig("5.1"))
    elements.append(fig("5.2"))
    elements.append(table(
        "5.1",
        "Traceability of contributions to source areas.",
        ["Contribution", "Representative source area", "Evaluation perspective"],
        [
            ["Whitelist-based enforcement", "agent/firewall, agent/whitelist, agent/network", "Does the endpoint apply central policy safely?"],
            ["Versioned synchronization", "server whitelist service/model and agent whitelist sync", "Does the agent avoid unnecessary policy reapplication?"],
            ["Teacher-scoped RBAC", "server RBAC config, group service, log service, whitelist service", "Does the backend restrict teacher visibility?"],
            ["Whitelist profiles", "whitelist profile controller, service, and model", "Can a lesson-specific policy be prepared and activated?"],
            ["Packet monitoring", "agent/capture and logging_module, server log controller/model", "Can observed traffic metadata be collected and reviewed?"],
        ],
    ))
    add_paragraphs(elements, [
        "The contributions are connected. Versioned synchronization improves firewall enforcement. RBAC protects whitelist profiles and logs. Packet monitoring gives feedback about enforced policy. These interactions are the main reason the project is treated as a distributed network security management system rather than as a single firewall script.",
        "This chapter has presented the main solutions and contributions of the thesis. The next chapter summarizes the outcome, compares the prototype with similar approaches, and describes future work needed before broader deployment."
    ])

    elements += [heading(1, "CHAPTER 6. CONCLUSION AND FUTURE WORK"), heading(2, "6.1 Conclusion")]
    add_paragraphs(elements, [
        "The thesis has designed and implemented SAINT, a distributed network security management system for educational computer laboratories. The system contains a Flask and MongoDB server, a server-rendered dashboard, REST APIs, SocketIO support, JWT and API-key authentication, RBAC, audit logging, and a Windows endpoint agent. The agent registers with the server, sends heartbeat information, synchronizes whitelist policy, resolves domains, manages firewall rules, captures packet metadata, extracts domains, and uploads logs.",
        "Compared with manual firewall configuration, SAINT provides central policy management and classroom workflow. Compared with DNS-only filtering, it binds policy to registered endpoint agents and group ownership. Compared with proxy-only approaches, it focuses on endpoint firewall enforcement and packet metadata rather than proxy configuration. Compared with commercial classroom suites, it is source-inspectable and specialized for the thesis problem, although it lacks the maturity and hardening of production products.",
        "The main achievements are the integrated design and implementation of server and agent components, approximately seventy API routes, twelve MongoDB collections, modular source packages, server-side tests for major backend behavior, static diagrams, and a reportable architecture. The main contributions are whitelist-based endpoint enforcement, versioned whitelist synchronization, teacher-scoped RBAC, whitelist profiles, and packet metadata logging.",
        "The work also reveals limitations. The prototype targets Windows endpoints and depends on local privileges for firewall operations. It does not implement full HTTPS content inspection, enterprise identity integration, signed auto-updates, large-scale performance testing, or advanced policy simulation. It should therefore be viewed as a strong educational prototype and research artifact, not yet as a production-ready security product.",
        "An important lesson from the project is that network security management is not only about blocking traffic. It requires identity, authorization, policy modeling, safe enforcement, monitoring, recovery, and maintainable software structure. A useful educational-lab system must let teachers express classroom intent while preserving administrator control and endpoint safety."
    ])
    elements += [heading(2, "6.2 Future work")]
    add_paragraphs(elements, [
        "The first future direction is operational safety. The agent should include an explicit dry-run mode, policy simulation, pre-apply validation, stronger restore workflows, emergency local disable mechanisms, and clearer status messages. Before enabling whitelist-only mode in a real laboratory, the system should prove that server access, DNS access, and required educational resources remain reachable.",
        "The second future direction is policy expressiveness. Whitelist profiles should support schedules, temporary overrides, exam mode, conflict detection, import validation, and approval workflow. Teachers should be able to preview which resources will be allowed before activating a profile. Administrators should be able to define global baseline rules that cannot be accidentally removed by group-level edits.",
        "The third future direction is deployment hardening. Production use should include TLS configuration, secure secret management, database backups, monitoring, centralized log retention, signed agent packages, versioned agent updates, and migration scripts. The server should provide health endpoints and observability metrics so that administrators can detect failures early.",
        "The fourth future direction is evaluation. The system should be tested with multiple endpoints, realistic traffic, intentional DNS failures, expired tokens, invalid whitelist entries, server downtime, and accidental firewall misconfiguration. Performance tests should measure synchronization latency, dashboard query response time, log ingestion throughput, and agent resource usage.",
        "The fifth future direction is privacy and compliance. Network logs should be minimized, retained for a defined duration, and presented only to authorized users. The system should document which metadata is collected and provide configuration for retention and export. This is important because educational environments involve student activity and institutional policies.",
        "With these improvements, SAINT can evolve from a graduation-thesis prototype into a more robust laboratory network management platform. The current source code provides a foundation for that path by separating server authority, dashboard workflow, endpoint enforcement, monitoring, and security controls."
    ])

    elements += [heading(1, "REFERENCE")]
    for ref in [
        "[1] L. L. Peterson and B. S. Davie, Computer Networks: A Systems Approach, 6th ed. Morgan Kaufmann, 2021.",
        "[2] M. Grinberg, Flask Web Development: Developing Web Applications with Python, 2nd ed. O'Reilly Media, 2018.",
        "[3] K. Chodorow, MongoDB: The Definitive Guide, 3rd ed. O'Reilly Media, 2019.",
        "[4] M. Jones, J. Bradley, and N. Sakimura, JSON Web Token (JWT), RFC 7519, IETF, 2015.",
        "[5] D. F. Ferraiolo and D. R. Kuhn, \"Role-Based Access Controls,\" in Proc. 15th National Computer Security Conference, 1992, pp. 554-563.",
        "[6] P. Biondi and the Scapy community, Scapy Documentation, https://scapy.readthedocs.io/, accessed May 26, 2026.",
        "[7] P. Mockapetris, Domain Names - Implementation and Specification, RFC 1035, IETF, 1987.",
        "[8] E. Rescorla, The Transport Layer Security (TLS) Protocol Version 1.3, RFC 8446, IETF, 2018.",
        "[9] D. Eastlake, Transport Layer Security (TLS) Extensions: Extension Definitions, RFC 6066, IETF, 2011.",
        "[10] R. Fielding and J. Reschke, Hypertext Transfer Protocol (HTTP/1.1): Message Syntax and Routing, RFC 7230, IETF, 2014.",
        "[11] N. Provos and D. Mazieres, \"A Future-Adaptable Password Scheme,\" in Proc. USENIX Annual Technical Conference, 1999, pp. 81-91.",
        "[12] Microsoft, \"netsh advfirewall firewall context,\" Microsoft Learn, https://learn.microsoft.com/windows-server/administration/windows-commands/netsh-advfirewall, accessed May 26, 2026.",
        "[13] Pallets Projects, Flask Documentation, https://flask.palletsprojects.com/, accessed May 26, 2026.",
        "[14] MongoDB Inc., MongoDB Manual, https://www.mongodb.com/docs/manual/, accessed May 26, 2026.",
        "[15] Qt Company, Qt for Python Documentation, https://doc.qt.io/qtforpython/, accessed May 26, 2026.",
        "[16] Python Software Foundation, Python Documentation, https://docs.python.org/3/, accessed May 26, 2026.",
    ]:
        elements.append(normal(ref))

    elements += [heading(1, "APPENDIX"), heading(1, "A. THESIS WRITING GUIDELINE")]
    add_paragraphs(elements, [
        "The official SOICT graduation-thesis template already contains a writing guideline appendix. In the final editable template, this appendix should be kept according to the school format. This generated draft includes a short compliance note only to document how the draft was prepared. The main chapters follow the template order: Introduction, Requirement Survey and Analysis, Theoretical Background and Technologies, Design, Implementation, and Evaluation, Solution and Contribution, and Conclusion and Future Work.",
        "The report avoids bullet-heavy writing in the main chapters and uses paragraphs for explanation. Tables and figures are referenced from body text. The abstract is written as one paragraph between 200 and 350 words. The acknowledgements section is between 100 and 150 words. References are listed in numbered IEEE-like format and exclude informal sources such as Wikipedia or lecture slides."
    ])
    elements += [heading(1, "B. USE CASE DESCRIPTIONS")]
    elements.append(table(
        "B.1",
        "Additional use case notes for implementation review.",
        ["Use case", "Implementation note", "Risk to verify"],
        [
            ["User login", "Implemented through dashboard authentication service and session/token state.", "Account lock, password hash, and logout behavior should be tested."],
            ["API-key creation", "API keys support agent enrollment and can be revoked.", "Key permissions and expiry should be verified before production use."],
            ["Agent heartbeat", "Agent sends status and metadata to the server.", "Offline detection threshold should match classroom expectations."],
            ["Whitelist import", "Bulk import supports faster policy preparation.", "Input validation should reject malformed domains or IP addresses."],
            ["Profile activation", "Active profile changes effective group policy.", "Activation conflicts should be handled clearly."],
            ["Log export", "Dashboard can export scoped log data.", "Export should respect RBAC and retention policy."],
            ["Firewall cleanup", "Agent should remove project-owned rules when appropriate.", "Cleanup must not delete unrelated user or system rules."],
        ],
    ))
    api_rows = extract_server_api_rows()
    if api_rows:
        elements += [heading(1, "C. SERVER API CATALOGUE")]
        add_paragraphs(elements, [
            "This appendix records the server API surface that was identified from the source-based report. The catalogue is included in the appendix rather than in the main body because the main thesis chapters should explain architecture and design decisions, while the full endpoint list is a technical reference. The route list is still important because it proves that the system contains a broad backend surface rather than a small demonstration script.",
            "The API catalogue can be read by group. Authentication routes support login, logout, refresh, verification, token information, and profile-related operations. User and audit routes support administrator accountability. Agent routes support enrollment, heartbeat, metadata update, group assignment, policy retrieval, and statistics. Whitelist routes support CRUD operations, import, export, bulk changes, statistics, and agent synchronization. Log routes support agent uploads, dashboard listing, clearing, export, and statistics. API-key routes support the lifecycle of enrollment credentials.",
            "The route list also shows the distinction between human-client endpoints and machine-client endpoints. Dashboard endpoints are protected by login sessions, role checks, or permission checks. Agent endpoints are protected by API keys during registration and JWT-based authentication after registration. This distinction is essential because the system should not grant a human dashboard session to an endpoint process and should not allow an endpoint token to perform administrative dashboard operations.",
            "In the final submission, the route catalogue can be shortened if the supervisor prefers a more compact report. However, keeping it in the appendix is useful for source traceability. It allows readers to map high-level functions from Chapter 2 and Chapter 4 to concrete HTTP routes in the implementation."
        ])
        elements.append(table(
            "C.1",
            "Server API route catalogue extracted from the source-based report.",
            ["Method", "Path", "Handler", "Authentication or RBAC", "Service or model"],
            api_rows,
        ))
    inventory = python_inventory_rows()
    if inventory:
        elements += [heading(1, "D. SOURCE MODULE INVENTORY")]
        add_paragraphs(elements, [
            "This appendix lists Python modules discovered through static AST parsing. The generator reads Python source files but does not import them. This matters because importing the agent or server could trigger runtime initialization, configuration loading, network access, or side effects. Static parsing is sufficient for a module inventory because it can identify top-level classes and functions without executing application code.",
            "The module inventory supports maintainability evaluation. A project with separate modules for controllers, services, models, middleware, GUI views, lifecycle, whitelist, firewall, capture, logging, and configuration is easier to inspect than a single script. The inventory also helps identify ownership boundaries. Server modules generally belong to web, business, persistence, and testing concerns, while agent modules generally belong to desktop UI, endpoint lifecycle, network policy, capture, and runtime state concerns.",
            "The inventory is not intended to replace source code. It is a report-level overview that helps readers understand the size and organization of the project. Detailed class diagrams can be produced from the same source if the final report needs more UML-style documentation."
        ])
        elements.append(table(
            "D.1",
            "Python module inventory extracted statically without importing agent or server runtime code.",
            ["Area", "Module", "Top-level classes", "Top-level functions"],
            inventory,
        ))
    elements += [heading(1, "E. API GROUP EXPLANATION")]
    add_paragraphs(elements, [
        "The authentication API group is the entry point for dashboard security. It includes login, logout, token refresh, token verification, token information, profile update, and password change behavior. The main design concern is separating authentication from authorization. Authentication proves who the user is, while authorization determines what that user can do. In SAINT, an authenticated teacher still needs group filtering before data is returned. This prevents a common implementation mistake where the application checks login state but forgets resource ownership.",
        "The user-management API group is administrator-oriented. It supports listing, creating, reading, updating, deleting, resetting passwords, and viewing statistics for dashboard accounts. This group is sensitive because account changes can affect all other security controls. For that reason, the implementation should use administrator checks, password hashing, audit records, and validation. A production version should also add stronger password policies, multi-factor authentication, and administrative approval workflows.",
        "The API-key group manages machine enrollment. API keys are not the same as dashboard passwords. They are credentials used by endpoint agents to prove that they are allowed to register. The design should store only hashes or protected forms of keys, support revocation, support expiration, and record usage metadata. This makes it possible to recover from a leaked enrollment key without changing all dashboard user passwords.",
        "The agent API group connects endpoint machines to the server. Registration creates or updates identity. Heartbeat updates online status and can indicate whether policy synchronization is needed. Agent detail and policy endpoints allow the dashboard to manage individual endpoints. The main design challenge is consistency: the server must know which group an agent belongs to, what policy applies to it, and when it last communicated.",
        "The group API group represents classroom organization. Groups connect teachers, agents, whitelist entries, profiles, and logs. A group is therefore more than a label. It is the unit of educational ownership used by RBAC filtering. Group APIs must be careful when assigning teachers or agents because an incorrect assignment can either hide needed data or expose data to the wrong teacher.",
        "The whitelist API group is the policy-management center. It includes basic operations, bulk operations, import/export, statistics, and the agent synchronization endpoint. The synchronization endpoint is special because it is consumed by agents rather than dashboard users. It must compute an effective policy from global entries, group entries, active profiles, and policy mode. This computation should be deterministic and version-aware so that agents can apply updates safely.",
        "The whitelist-profile API group adds classroom workflow on top of base whitelist entries. Profiles let a teacher prepare a lesson-specific policy and activate it for a group. This avoids repeated manual editing of the base whitelist. The design should prevent conflicting active profiles and should record which user activated a profile. In future work, profiles can support schedules, exam mode, and approval workflow.",
        "The log API group supports observation and review. Agents upload batches of network metadata, and dashboard users query logs, statistics, exports, or clear operations. Log routes must respect RBAC because logs can reveal student activity or class-specific behavior. The project therefore treats log access as a security-sensitive operation rather than a simple data table.",
        "The audit API group supports accountability. Audit logs are different from network logs because they record management actions rather than endpoint traffic. Audit records help answer questions such as who changed a whitelist, who revoked an API key, or who updated a user account. This distinction is important in a system that can change endpoint network behavior."
    ])
    elements += [heading(1, "F. OPERATIONAL SAFETY AND DEPLOYMENT PLAN")]
    add_paragraphs(elements, [
        "Operational safety is especially important for SAINT because the agent can enable default-deny firewall behavior. A normal web application bug may cause an error message or failed request, but a firewall bug can prevent a machine from reaching the server or the Internet. For this reason, deployment must be treated as a staged process. The report-generation process intentionally does not run the agent, the server, Docker compose, packet capture, netsh, or firewall code.",
        "The first deployment stage should be source and configuration review. Administrators should review .env-example, required environment variables, server URL format, allowed DNS servers, and default policy mode. Real secrets should be created outside the repository. The server should be configured with strong JWT secrets, secure database credentials, and TLS termination. The agent should be configured with a test server URL and a controlled enrollment API key.",
        "The second stage should be server-only testing. The server can be tested with a temporary database, seeded users, seeded groups, and test API keys. Administrators should verify login, role separation, group filtering, whitelist operations, profile activation, log routes, and audit records. During this stage, no endpoint firewall behavior is enabled. The goal is to prove that the policy authority behaves correctly before any endpoint consumes policy.",
        "The third stage should be agent dry-run testing. A small number of non-critical Windows machines should run the agent without enabling whitelist-only enforcement. The test should verify registration, heartbeat, token refresh, whitelist download, DNS resolution, status display, and log upload. If the agent supports a dry-run or observe-only mode, that mode should be used first. The purpose is to validate communication and visibility before allowing firewall changes.",
        "The fourth stage should be limited enforcement testing. One test group and one or two endpoints should be selected. The whitelist should include the server address, DNS resolver, required Windows services if needed, and a small set of learning resources. Administrators should confirm that the agent can still reach the server after policy application. A local recovery account and manual firewall restore procedure should be prepared before the test starts.",
        "The fifth stage should be classroom pilot operation. A teacher should prepare a real lesson profile, activate it for a small group, and observe whether students can access required resources. Logs should be reviewed after the session to identify missing whitelist entries or unexpected traffic. The pilot should be considered successful only if policy activation, synchronization, enforcement, logging, and recovery procedures are all understood by the operator.",
        "The final stage is broader rollout. At this stage, administrators should define retention policies, backup schedules, agent update procedures, monitoring dashboards, incident response steps, and user training material. The deployment should still allow quick rollback. The system should never depend on a single teacher knowing how to recover blocked network access.",
    ])
    elements.append(table(
        "F.1",
        "Recommended deployment stages for safer rollout.",
        ["Stage", "Main activity", "Exit condition"],
        [
            ["Configuration review", "Prepare secrets, server URL, API keys, TLS, and database settings", "No real secret is stored in source and all required variables are documented"],
            ["Server-only test", "Run backend tests and validate dashboard workflows with test data", "Authentication, RBAC, whitelist, logs, groups, and audit behave as expected"],
            ["Agent dry-run", "Register endpoint and verify communication without enforcement", "Heartbeat, token refresh, sync, DNS resolution, and log upload work"],
            ["Limited enforcement", "Enable whitelist-only mode on one or two test machines", "Server, DNS, and learning resources remain reachable and rollback works"],
            ["Classroom pilot", "Use a lesson profile with a small group", "Teacher can activate policy and review logs without IT intervention"],
            ["Broader rollout", "Expand to more machines with monitoring and backup", "Support team can operate and recover the system consistently"],
        ],
    ))
    elements += [heading(1, "G. LIMITATIONS AND RISK REGISTER")]
    add_paragraphs(elements, [
        "A thesis report should not only describe what works; it should also document known limitations. The first limitation is platform scope. The agent targets Windows endpoints and depends on Windows-specific integration for firewall behavior. This is appropriate for many computer laboratories, but it does not cover Linux, macOS, tablets, or browser-only environments. Cross-platform support would require a different enforcement abstraction.",
        "The second limitation is network complexity. Modern websites use content delivery networks, multiple IP addresses, encrypted protocols, DNS caching, and sometimes privacy-preserving name resolution. A domain whitelist may resolve to many addresses, and those addresses may change. The agent must therefore refresh resolution results and avoid assuming that a single DNS lookup is permanent. Some traffic may not expose useful domain metadata through DNS, HTTP Host, or TLS SNI.",
        "The third limitation is policy completeness. A teacher may forget to include a required dependency such as a package repository, authentication domain, CDN domain, certificate validation endpoint, or learning-platform subdomain. In whitelist-only mode, this omission can interrupt the lesson. Policy simulation and profile templates would reduce this risk, but they are not complete in the current prototype.",
        "The fourth limitation is privilege and trust. A firewall-enforcing agent requires local privileges and must be trusted by the endpoint. If the agent is misconfigured, tampered with, or stopped, enforcement may fail. A production system should include service-mode installation, integrity protection, signed updates, and tamper-resistant configuration.",
        "The fifth limitation is scalability. The source structure supports a useful prototype, but broad deployment would require load testing, database retention planning, log ingestion benchmarks, and monitoring. Logs can grow quickly if packet capture is verbose. Indexes and retention policies must be tuned before a large laboratory rollout.",
        "The sixth limitation is privacy. Network metadata can reveal user behavior. Even when content is not decrypted, domains and timestamps can be sensitive. The system should define who can see logs, how long logs are retained, how exports are protected, and how students are informed about monitoring policies. RBAC is a necessary foundation but not a complete privacy policy.",
        "The seventh limitation is template integration. This generated report is a DOCX draft based on the PDF template outline. If the school provides an editable Word or LaTeX template, the content should be copied into that official file so that page styles, captions, page numbering, and the fixed Appendix A match school requirements exactly."
    ])
    elements.append(table(
        "G.1",
        "Risk register for the SAINT prototype.",
        ["Risk", "Impact", "Mitigation in current design", "Future improvement"],
        [
            ["Default-deny rule blocks connectivity", "Agent or student machine loses network access", "Self-allow, server allow, DNS allow, snapshot, restore, cleanup", "Dry-run mode and automatic rollback"],
            ["Teacher sees unrelated data", "Privacy and operational boundary violation", "RBAC and group filtering", "More systematic authorization tests"],
            ["API key is leaked", "Unauthorized agent enrollment", "API-key service, validation, revocation", "Short-lived enrollment windows and approval workflow"],
            ["Whitelist misses required dependency", "Lesson interruption", "Profiles and import/export support", "Policy simulation and dependency suggestions"],
            ["Logs grow too quickly", "Storage and query performance issues", "Indexes and log models", "Retention policies and aggregation pipeline"],
            ["Agent package is tampered with", "Endpoint trust is broken", "Source-controlled packaging spec", "Signed updates and integrity verification"],
            ["Server secret is misconfigured", "Token or database compromise", ".env-example separates configuration", "Secret manager integration and deployment checklist"],
        ],
    ))
    elements += [heading(1, "H. EVALUATION PLAN FOR FINAL DEFENSE")]
    add_paragraphs(elements, [
        "The current report is based on static source analysis and document generation. For the final defense, the project should be evaluated through a controlled experiment that does not risk the student's primary network connection. The purpose of the evaluation is to demonstrate that the architecture described in the thesis is not only a code organization but also a working security-management workflow. The evaluation should be prepared in a separate laboratory network or virtualized test environment whenever possible.",
        "The first evaluation scenario is administrator setup. The evaluator creates or verifies administrator access, creates teacher accounts, creates groups, creates API keys, imports or adds whitelist entries, and checks audit logs. The expected result is that all administrative actions succeed and are recorded where appropriate. This scenario demonstrates dashboard completeness, persistence, and accountability.",
        "The second evaluation scenario is teacher group operation. A teacher logs in and views only assigned groups. The teacher prepares a whitelist profile for a lesson, activates it, reviews scoped agents, and inspects scoped logs. The expected result is that the teacher cannot access unrelated groups or logs. This scenario demonstrates that RBAC and group filtering are not only theoretical claims but visible behavior.",
        "The third evaluation scenario is agent enrollment. A test endpoint agent uses an API key to register with the server and receives identity and token material. The server displays the agent, records last-seen status after heartbeat, and associates the agent with a group. The expected result is that the dashboard reflects endpoint identity and runtime state. This scenario demonstrates the machine-client side of the architecture.",
        "The fourth evaluation scenario is whitelist synchronization without enforcement. The agent pulls the effective whitelist and reports synchronization status while firewall enforcement remains disabled or in dry-run mode. The expected result is that the agent receives the correct global, group, and profile-derived entries and does not reapply policy when versions are unchanged. This scenario validates synchronization logic without risking connectivity.",
        "The fifth evaluation scenario is limited enforcement. A carefully prepared whitelist is applied on a disposable or test endpoint. The whitelist must include the server, DNS, and required learning resources. The evaluator checks that allowed resources remain reachable and non-whitelisted resources are blocked when whitelist-only mode is active. The expected result is correct local enforcement and successful restoration after the test. This scenario should be performed only after the previous scenarios pass.",
        "The sixth evaluation scenario is log collection. The endpoint generates controlled traffic to allowed and disallowed destinations. The agent captures metadata and uploads logs. The dashboard shows logs and statistics according to role. The expected result is that log entries include useful metadata such as timestamp, agent identity, address, protocol, and domain when extractable. This scenario validates the monitoring contribution.",
        "The seventh evaluation scenario is failure handling. The evaluator tests expired token behavior, unavailable server behavior, invalid API key behavior, invalid whitelist entry behavior, DNS resolution failure, and policy rollback behavior. The expected result is that the system fails visibly and safely. A security management system should not hide errors that affect policy enforcement.",
        "The evaluation should record both successful and unsuccessful outcomes. If a test fails, the report should explain whether the failure is a bug, a known limitation, a configuration problem, or an out-of-scope feature. This is important because the graduation thesis is evaluated not only by positive demonstrations but also by the student's ability to reason about limitations and engineering trade-offs."
    ])
    elements.append(table(
        "H.1",
        "Suggested evaluation scenarios for the final defense.",
        ["Scenario", "Observed function", "Expected evidence"],
        [
            ["Administrator setup", "Users, groups, API keys, whitelist, audit", "Dashboard operations and audit records"],
            ["Teacher operation", "RBAC and group ownership", "Teacher sees only assigned groups and logs"],
            ["Agent enrollment", "API-key registration and heartbeat", "Agent appears online in dashboard"],
            ["Dry-run synchronization", "Effective whitelist and version metadata", "Agent receives expected policy without enforcement"],
            ["Limited enforcement", "Windows Firewall allow/default-deny behavior", "Allowed resources work and blocked resources fail"],
            ["Log collection", "Packet metadata capture and upload", "Dashboard shows scoped log records and statistics"],
            ["Failure handling", "Token, API key, DNS, server, rollback behavior", "Errors are visible and recovery path is documented"],
        ],
    ))
    elements += [heading(1, "I. FINAL EDITING CHECKLIST")]
    add_paragraphs(elements, [
        "Before submitting the thesis, the content should be copied into the official editable template if the school provides one. The PDF template used for this task is not directly editable. The generated DOCX follows the template outline, but final formatting such as page numbers, caption numbering, section breaks, and institutional front-page details should be checked in Microsoft Word or the official LaTeX source.",
        "The first editing task is metadata completion. The student name and student ID are already filled with Bui Xuan Son and 20225586. The supervisor name, student email, and program name remain placeholders because they were not reliably present in the project source. These fields must be completed before submission. The cover page and signature page should match the school template exactly.",
        "The second editing task is table-of-contents update. The DOCX contains a Word TOC field. After opening the document in Word, the student should update the table of contents and check that Chapter 1 begins after the front matter. The lists of figures and tables are generated as text lists and may be converted to Word fields if strict formatting is required.",
        "The third editing task is figure review. The figures in this draft are source-derived static diagrams, not runtime screenshots. They are appropriate for architecture and design explanation. If the supervisor requires real UI screenshots, those screenshots should be captured in a safe environment that does not enable accidental network blocking. Screenshots should not contain secrets, database credentials, or private student data.",
        "The fourth editing task is language consistency. The main report is written in English because the template is an English thesis template. Vietnamese notes from the original school template should not remain in the final body unless they are part of the unchanged official appendix. Technical terms should be used consistently: Administrator, Teacher, Windows Agent, Server, Dashboard, Whitelist, Group, Profile, RBAC, JWT, API key, and MongoDB collection.",
        "The fifth editing task is source alignment. If the code changes after this report is generated, the script should be rerun and the changed sections should be reviewed. The most likely sections to become outdated are API catalogue, module inventory, technology table, database collections, and deployment artifacts. The thesis should not claim features that are no longer present in source code.",
        "The sixth editing task is final evidence insertion. The final report may include additional test results, screenshots, performance measurements, or deployment logs. These should be inserted into Chapter 4 or Appendix H depending on their importance. Evidence should be concise and should support the thesis claims rather than simply increasing page count.",
        "The seventh editing task is reference verification. Internet references should be official sources and should include access dates. Informal blogs, Wikipedia pages, lecture slides, and unverified tutorials should be avoided. The current reference list uses books, RFCs, official documentation, and a classic RBAC paper, which is consistent with the school guideline.",
        "The final editing task is safety review. Any instruction that could run the agent, change firewall rules, start packet capture, or run network-affecting commands should be clearly separated from normal document editing. Report generation must remain a static process. Practical demonstrations should be done only in an isolated test environment with a recovery plan.",
        "After these edits are complete, the student should export a final PDF from the official template and compare the exported page order against the school checklist. This final pass should also check that every table and figure is referenced from the surrounding text."
    ])
    elements.append(table(
        "I.1",
        "Final thesis editing checklist.",
        ["Checklist item", "Action required before submission"],
        [
            ["Supervisor and program metadata", "Replace placeholders on both cover pages"],
            ["Table of contents", "Open in Word and update the TOC field"],
            ["Lists of figures and tables", "Check numbering and formatting after edits"],
            ["Figures", "Replace or supplement diagrams with safe screenshots if required"],
            ["Source alignment", "Rerun generator after code changes and review affected sections"],
            ["Test evidence", "Add controlled test results from final evaluation"],
            ["References", "Verify official sources and access dates"],
            ["Safety", "Do not run firewall-affecting agent behavior outside a prepared environment"],
        ],
    ))
    return rebalance_main_chapters(elements)


def set_table_header(row) -> None:
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "B5121B")
        tc_pr.append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    for style_name, size in (("Title", 18), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        if "Heading" in style_name:
            style.font.color.rgb = RGBColor(181, 18, 27)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def add_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    p = doc.add_paragraph("HANOI UNIVERSITY OF SCIENCE AND TECHNOLOGY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p = doc.add_paragraph("SCHOOL OF INFORMATION AND COMMUNICATIONS TECHNOLOGY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    doc.add_paragraph()
    p = doc.add_paragraph("GRADUATION THESIS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(18)
    doc.add_paragraph()
    p = doc.add_paragraph(TITLE, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(181, 18, 27)
    doc.add_paragraph()
    for line in [
        f"Student: {STUDENT}",
        f"Student ID: {STUDENT_ID}",
        f"Email: {EMAIL}",
        f"Program: {PROGRAM}",
        f"Supervisor: {SUPERVISOR}",
        f"Department: {DEPARTMENT}",
        f"School: {SCHOOL}",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph("HANOI, 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    p = doc.add_paragraph("HANOI UNIVERSITY OF SCIENCE AND TECHNOLOGY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p = doc.add_paragraph("SCHOOL OF INFORMATION AND COMMUNICATIONS TECHNOLOGY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    doc.add_paragraph()
    p = doc.add_paragraph("GRADUATION THESIS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(18)
    doc.add_paragraph()
    p = doc.add_paragraph(TITLE, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(181, 18, 27)
    doc.add_paragraph()
    for line in [
        f"Student: {STUDENT}",
        f"Student ID: {STUDENT_ID}",
        f"Email: {EMAIL}",
        f"Program: {PROGRAM}",
        f"Supervisor: {SUPERVISOR}",
        "",
        "Signature",
        "",
        f"Department: {DEPARTMENT}",
        f"School: {SCHOOL}",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("HANOI, 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_front_matter(doc: Document, elements: list[dict]) -> None:
    doc.add_heading("Acknowledgements", level=1)
    doc.add_paragraph(ACKNOWLEDGEMENTS)
    doc.add_page_break()
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(ABSTRACT)
    doc.add_page_break()
    add_toc(doc)
    doc.add_heading("List of Figures", level=1)
    for figure in [element["figure"] for element in elements if element["type"] == "figure"]:
        doc.add_paragraph(f"Figure {figure.number}. {figure.title}")
    doc.add_page_break()
    doc.add_heading("List of Tables", level=1)
    for element in elements:
        if element["type"] == "table":
            doc.add_paragraph(f"Table {element['number']}. {element['title']}")
    doc.add_page_break()
    doc.add_heading("List of Abbreviations", level=1)
    rows = [
        ["SAINT", "Security Agent Integrated Network Tool"],
        ["API", "Application Programming Interface"],
        ["CLI", "Command Line Interface"],
        ["CRUD", "Create, Read, Update, Delete"],
        ["DNS", "Domain Name System"],
        ["GUI", "Graphical User Interface"],
        ["HMAC", "Hash-based Message Authentication Code"],
        ["HTTP", "Hypertext Transfer Protocol"],
        ["JWT", "JSON Web Token"],
        ["LRU", "Least Recently Used"],
        ["MVC", "Model-View-Controller"],
        ["MVP", "Model-View-Presenter"],
        ["RBAC", "Role-Based Access Control"],
        ["REST", "Representational State Transfer"],
        ["SNI", "Server Name Indication"],
        ["TLS", "Transport Layer Security"],
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Abbreviation"
    t.rows[0].cells[1].text = "Meaning"
    set_table_header(t.rows[0])
    for abbr, meaning in rows:
        cells = t.add_row().cells
        cells[0].text = abbr
        cells[1].text = meaning
    doc.add_page_break()


def add_element_docx(doc: Document, element: dict) -> None:
    if element["type"] == "h":
        doc.add_heading(element["text"], level=element["level"])
    elif element["type"] == "p":
        p = doc.add_paragraph(element["text"])
        p.paragraph_format.first_line_indent = Inches(0.25)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif element["type"] == "table":
        p = doc.add_paragraph(f"Table {element['number']}. {element['title']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        headers = element["headers"]
        rows = element["rows"]
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for idx, header in enumerate(headers):
            t.rows[0].cells[idx].text = header
        set_table_header(t.rows[0])
        for row in rows:
            cells = t.add_row().cells
            for idx, value in enumerate(row):
                cells[idx].text = value
        doc.add_paragraph()
    elif element["type"] == "figure":
        figure: Figure = element["figure"]
        image_path = ASSET_DIR / figure.filename
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.2))
        c = doc.add_paragraph(f"Figure {figure.number}. {figure.title}")
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].bold = True
        doc.add_paragraph()


def build_docx(elements: list[dict]) -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_front_matter(doc, elements)
    for section in doc.sections:
        add_page_number(section)
    for element in elements:
        add_element_docx(doc, element)
    doc.save(OUT_DOCX)


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        safe = [cell.replace("\n", " ").replace("|", "\\|") for cell in row]
        lines.append("| " + " | ".join(safe) + " |")
    return "\n".join(lines)


def build_markdown(elements: list[dict]) -> None:
    lines: list[str] = []
    lines += [
        f"# {TITLE}",
        "",
        f"**Student:** {STUDENT}",
        f"**Student ID:** {STUDENT_ID}",
        f"**Email:** {EMAIL}",
        f"**Program:** {PROGRAM}",
        f"**Supervisor:** {SUPERVISOR}",
        f"**Department:** {DEPARTMENT}",
        f"**School:** {SCHOOL}",
        "",
        "## Acknowledgements",
        "",
        ACKNOWLEDGEMENTS,
        "",
        "## Abstract",
        "",
        ABSTRACT,
        "",
        "## Table of Contents",
        "",
        "The DOCX version contains a Word table-of-contents field. In Word, use **Update Field** to refresh page numbers after final editing.",
        "",
        "## List of Figures",
        "",
    ]
    for element in elements:
        if element["type"] == "figure":
            figure: Figure = element["figure"]
            lines.append(f"Figure {figure.number}. {figure.title}")
    lines += [
        "",
        "## List of Tables",
        "",
    ]
    for element in elements:
        if element["type"] == "table":
            lines.append(f"Table {element['number']}. {element['title']}")
    lines += [
        "",
        "## List of Abbreviations",
        "",
        md_table(
            ["Abbreviation", "Meaning"],
            [
                ["SAINT", "Security Agent Integrated Network Tool"],
                ["API", "Application Programming Interface"],
                ["CLI", "Command Line Interface"],
                ["CRUD", "Create, Read, Update, Delete"],
                ["DNS", "Domain Name System"],
                ["GUI", "Graphical User Interface"],
                ["HMAC", "Hash-based Message Authentication Code"],
                ["HTTP", "Hypertext Transfer Protocol"],
                ["JWT", "JSON Web Token"],
                ["LRU", "Least Recently Used"],
                ["MVC", "Model-View-Controller"],
                ["MVP", "Model-View-Presenter"],
                ["RBAC", "Role-Based Access Control"],
                ["REST", "Representational State Transfer"],
                ["SNI", "Server Name Indication"],
                ["TLS", "Transport Layer Security"],
            ],
        ),
        "",
    ]
    for element in elements:
        if element["type"] == "h":
            lines.append("#" * element["level"] + " " + element["text"])
            lines.append("")
        elif element["type"] == "p":
            lines.append(element["text"])
            lines.append("")
        elif element["type"] == "table":
            lines.append(f"**Table {element['number']}. {element['title']}**")
            lines.append("")
            lines.append(md_table(element["headers"], element["rows"]))
            lines.append("")
        elif element["type"] == "figure":
            figure: Figure = element["figure"]
            asset_path = Path("report_assets") / figure.filename
            lines.append(f"![Figure {figure.number}. {figure.title}]({asset_path.as_posix()})")
            lines.append("")
            lines.append(f"**Figure {figure.number}. {figure.title}**")
            lines.append("")
    OUT_MD.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def validate_outputs() -> None:
    if not OUT_DOCX.exists() or OUT_DOCX.stat().st_size < 10_000:
        raise RuntimeError("DOCX output was not generated correctly")
    if not OUT_MD.exists() or OUT_MD.stat().st_size < 10_000:
        raise RuntimeError("Markdown output was not generated correctly")
    with zipfile.ZipFile(OUT_DOCX) as archive:
        names = set(archive.namelist())
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise RuntimeError("DOCX package is invalid")
    md_text = OUT_MD.read_text(encoding="utf-8")
    forbidden = ["MONGO_URI=", "JWT_SECRET", "SECRET_KEY=", "PASSWORD=", "API_KEY="]
    leaked = [token for token in forbidden if token in md_text]
    if leaked:
        raise RuntimeError(f"Potential secret tokens found in report: {leaked}")
    for phrase in ("TODO", "Lorem ipsum"):
        if phrase in md_text:
            raise RuntimeError(f"Placeholder phrase found: {phrase}")
    ack_words = word_count(ACKNOWLEDGEMENTS)
    abstract_words = word_count(ABSTRACT)
    if not (100 <= ack_words <= 150):
        raise RuntimeError(f"Acknowledgements word count out of range: {ack_words}")
    if not (200 <= abstract_words <= 350):
        raise RuntimeError(f"Abstract word count out of range: {abstract_words}")
    references = re.findall(r"^\[\d+\]", md_text, flags=re.MULTILINE)
    if len(references) < 12:
        raise RuntimeError("Not enough IEEE-style references")
    total_words = word_count(md_text)
    if total_words < 18_000:
        raise RuntimeError(f"Expanded report is still too short: {total_words} words")


def main() -> None:
    create_diagram_assets()
    elements = build_long_elements()
    build_docx(elements)
    build_markdown(elements)
    validate_outputs()
    print(f"Generated: {OUT_DOCX}")
    print(f"Generated: {OUT_MD}")
    print(f"Assets: {ASSET_DIR}")
    print(f"Acknowledgements words: {word_count(ACKNOWLEDGEMENTS)}")
    print(f"Abstract words: {word_count(ABSTRACT)}")
    print(f"Markdown words: {word_count(OUT_MD.read_text(encoding='utf-8'))}")


if __name__ == "__main__":
    main()
